#!/usr/bin/env python3
"""
Reproduces two bugs reported against the real case_smith (17-document) Case Mode run:

1. The exported chronology's PAGE column has no Bates numbers.
2. By the time processing finishes, the PDF viewer says the source documents "are no longer on
   disk", even though they genuinely are.

Root cause found by reading the code (see app.js's fetchAndRenderCaseResult): a browser tab that
opens the review page WHILE a case is still processing polls for live updates and only re-renders
when the finding count GROWS. But the final "done" result runs Bates resolution AND duplicate-record
merging (which can REDUCE the count) after all chunks are already merged — so a tab left open
across that transition could silently keep showing the last pre-resolution snapshot forever (no
Bates numbers, and polling stops once the job is "done" regardless of whether that final render
ever happened). This test drives the EXACT real scenario: start a real case_smith run, open the
review page immediately (before it's finished, simulating a reviewer who starts watching right
away), leave it open across the entire run, and check the page correctly updates to the final,
Bates-resolved result once done -- without ever reloading the tab.

This is a slow test (case_smith is a real 17-document case, ~15-25 minutes end to end). Requires
the test server already running on http://127.0.0.1:5051 (see start_test_server.sh).
Usage: testing/.venv/bin/python testing/verify_case_smith_live_view.py [--headed]
"""
import sys
import time
import urllib.request
import json
from pathlib import Path

from docx import Document
from playwright.sync_api import sync_playwright

from case_upload_helpers import urllib_start_case

APP_URL = "http://127.0.0.1:5051"
BASE = Path(__file__).resolve().parent.parent
CASE_FOLDER = BASE / "sample_data" / "case_smith"
MAX_WAIT_S = 30 * 60  # case_smith is large; give it generous room


def start_case() -> str:
    return urllib_start_case(
        APP_URL, "Jane Smith LiveView Repro", CASE_FOLDER, priority_hint="records", timeout=120,
    )


def get_manifest(job_id: str) -> dict:
    with urllib.request.urlopen(f"{APP_URL}/case/status/{job_id}", timeout=15) as resp:
        return json.loads(resp.read().decode("utf-8"))


def primary_group_key(manifest: dict) -> str:
    # Mirrors _normalize_patient_key's slugification closely enough for this fixture's plain-ASCII
    # name (lowercase, strip non-alphanumerics) -- confirmed against the actual manifest below
    # rather than trusted blindly.
    return "".join(ch for ch in manifest["plaintiff_name"].lower() if ch.isalnum())


def main():
    headed = "--headed" in sys.argv
    if not CASE_FOLDER.is_dir():
        print(f"FAIL: {CASE_FOLDER} does not exist on this machine — can't run this test")
        return 1

    job_id = start_case()
    print(f"Started real case_smith run: job_id={job_id}")
    manifest = get_manifest(job_id)
    group_key = next(iter(manifest.get("groups", {})), None) or primary_group_key(manifest)
    print(f"primary group key: {group_key}")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=not headed)
        try:
            ok = run_checks(browser, job_id, group_key)
        finally:
            browser.close()
    print("OVERALL:", "PASS" if ok else "FAIL")
    return 0 if ok else 1


def run_checks(browser, job_id: str, group_key: str) -> bool:
    all_ok = True
    page = browser.new_page(viewport={"width": 1400, "height": 1000})
    console_errors = []
    page.on("pageerror", lambda e: console_errors.append(f"[PAGEERROR] {e}"))
    page.on("console", lambda m: console_errors.append(f"[console.error] {m.text}") if m.type == "error" else None)

    # Open the review page RIGHT AWAY, before the job is anywhere near done -- this is the exact
    # scenario that triggers the bug (a tab that's been live-polling since before "done").
    url = f"{APP_URL}/review?case={job_id}&group={group_key}"
    print(f"Opening {url} immediately (job is still processing)...")
    page.goto(url)
    page.wait_for_timeout(2000)

    # Now poll the manifest directly (not through the browser) until the job is genuinely done,
    # while the ALREADY-OPEN tab keeps doing its own live-polling in the background exactly as a
    # real user's tab would.
    start = time.time()
    manifest = None
    while time.time() - start < MAX_WAIT_S:
        manifest = get_manifest(job_id)
        status = manifest["groups"].get(group_key, {}).get("status")
        elapsed = int(time.time() - start)
        print(f"  [{elapsed}s] job status: {manifest['status']}, group status: {status}, "
              f"progress: {manifest['groups'].get(group_key, {}).get('progress_text')}")
        if manifest["status"] == "done":
            break
        time.sleep(15)
    else:
        print(f"FAIL: case_smith did not finish within {MAX_WAIT_S}s")
        return False

    print(f"\nJob finished after {int(time.time() - start)}s of waiting. "
          f"Now checking the ALREADY-OPEN tab (never reloaded) updated correctly...\n")

    # Give the tab's own polling loop (every 4s) a couple of cycles to catch up.
    page.wait_for_timeout(10_000)

    # --- Check 1: the citation lines shown in the still-open tab now include real Bates numbers,
    # not "Bates not resolved" -- proving the final, Bates-resolved render actually happened. ---
    source_lines = page.locator(".finding-source").all_inner_texts()
    print(f"'.finding-source' lines in the still-open tab ({len(source_lines)} total):")
    for line in source_lines[:15]:
        print(f"  {line!r}")
    resolved_count = sum(1 for s in source_lines if "Bates" in s and "not resolved" not in s)
    ok1 = resolved_count > 0
    print(f"-> {resolved_count}/{len(source_lines)} show a real resolved Bates number "
          f"=> {'PASS' if ok1 else 'FAIL'} (the live-open tab must show the final resolved result)\n")
    all_ok &= ok1

    # --- Check 2: clicking a finding's citation actually loads the PDF (no "session expired" /
    # "no longer on disk" error) in this SAME tab that's been open the whole time. ---
    citation = page.locator(".finding-citation").first
    if citation.count():
        citation.click()
        page.wait_for_timeout(1500)
        viewer_title = page.locator("#viewer-title").inner_text()
        viewer_hint = page.locator("#viewer-hint").inner_text()
        highlight_count = page.locator(".pdf-highlight").count()
        print(f"viewer-title: {viewer_title!r}")
        print(f"viewer-hint: {viewer_hint!r}")
        print(f"highlight boxes drawn: {highlight_count}")
        pdf_ok = "expired" not in viewer_title.lower() and "no longer on disk" not in viewer_hint.lower()
        print(f"-> {'PASS' if pdf_ok else 'FAIL'} (source PDF must actually load, not '404/expired')\n")
        all_ok &= pdf_ok
    else:
        print("No citation chip found to test PDF loading against.\n")
        pdf_ok = False
        all_ok = False

    # --- Check 3: approve everything and confirm the exported .docx PAGE column has real Bates
    # numbers, not blank -- the original, concretely-reported symptom. ---
    findings_count = page.locator(".finding").count()
    for i in range(findings_count):
        f = page.locator(".finding").nth(i)
        if "status-approved" not in (f.get_attribute("class") or ""):
            f.click()
            page.wait_for_timeout(100)
            page.keyboard.press("Escape")  # back out of the click-opens-edit-mode textarea
            page.wait_for_timeout(100)
            page.locator("#approve-btn").click()
            page.wait_for_timeout(100)

    with page.expect_download() as dl_info:
        page.locator("#export-btn").click()
        page.wait_for_timeout(300)
        proceed_btn = page.locator("#export-confirm-proceed")
        if proceed_btn.is_visible():
            proceed_btn.click()
    download = dl_info.value
    export_path = Path("/tmp/lfa_case_smith_export_test.docx")
    download.save_as(str(export_path))
    doc = Document(export_path)
    table = doc.tables[0] if doc.tables else None
    page_values = []
    if table:
        for row in table.rows[1:]:
            page_values.append(row.cells[1].text.strip())
    non_blank_pages = [v for v in page_values if v and v.lower() != "(not provided)"]
    print(f"PAGE column values in exported .docx: {page_values}")
    ok3 = len(non_blank_pages) > 0
    print(f"-> {len(non_blank_pages)}/{len(page_values)} rows have a real Bates/page value "
          f"=> {'PASS' if ok3 else 'FAIL'}\n")
    all_ok &= ok3

    if console_errors:
        print(f"Console errors seen during run ({len(console_errors)}):")
        for e in console_errors[:10]:
            print(f"  - {e}")
        all_ok = False
    else:
        print("No console errors during run.")

    return all_ok


if __name__ == "__main__":
    sys.exit(main())
