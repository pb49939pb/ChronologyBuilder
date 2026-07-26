#!/usr/bin/env python3
"""Confirms the "Export Approved Chronology" button produces a real Microsoft Word document (not
a PDF or Markdown — see TEST_RESULTS.md for why this changed) matching the firm's own blank
chronology template structure (see ~/Downloads/Blank Chronology Template.docx) — not just "is a
.docx", but has the actual required sections: RE/Updated/DOL/Facts/DOB/Address/Social-Hx/PMHx/FHx/
Surgical-Hx/PCP header block, Abbreviations/Record Sources, Outstanding Records placeholder, a
DATE/PAGE/RECORD/SOURCE/DESCRIPTION table, and Times New Roman 12pt body text throughout (the
firm's explicit formatting requirement).

Uses a plain single-upload session (no Case Mode plaintiff/complaint) on purpose — this is exactly
the "no case metadata known" scenario, and the template's blank-placeholder behavior needs to work
here too, not just when Case Mode's plaintiff/defendant/DOL/facts are available.
"""
import sys
from pathlib import Path

from docx import Document
from playwright.sync_api import sync_playwright

BASE = Path(__file__).resolve().parent.parent
ZIP_PATH = BASE / "sample_data" / "case_000_pdfs.zip"
APP_URL = "http://127.0.0.1:5051"

REQUIRED_LABELS = [
    "COMBINED SUMMARY OF MEDICAL RECORDS / TIMELINE",
    "RE:", "Updated:", "DOL:", "Facts:", "DOB:", "ADDRESS:",
    "Social Hx:", "PMHx:", "FHx:", "SURGICAL HX:", "PCP:",
    "ABBREVIATIONS AND RECORD SOURCES",
    "Other record sources/Outstanding records:",
    "Chronological Events:",
]


def main():
    out_path = Path("/tmp/lfa_export_template_test.docx")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            page = browser.new_page()
            page.goto(f"{APP_URL}/review")
            page.set_input_files("#zipfile", str(ZIP_PATH))
            page.click("#submit-btn")
            page.wait_for_selector("#results", state="visible", timeout=120_000)
            page.wait_for_selector(".finding", timeout=10_000)

            # Approve one finding so export isn't blocked.
            page.locator(".finding").first.locator(".finding-text").click()
            page.wait_for_timeout(200)
            page.keyboard.press("Escape")  # exit edit mode, keep selection
            page.wait_for_timeout(200)
            page.keyboard.press("ArrowRight")  # approve
            page.wait_for_timeout(300)

            with page.expect_download(timeout=30_000) as dl_info:
                page.click("#export-btn")
                # This fixture can have more than one finding — if only one got approved, the
                # rest are still unreviewed, which raises a styled in-app modal (not a native
                # confirm() dialog, see #export-confirm-modal) asking to export anyway.
                page.wait_for_timeout(300)
                proceed_btn = page.locator("#export-confirm-proceed")
                if proceed_btn.is_visible():
                    proceed_btn.click()
            download = dl_info.value
            download.save_as(str(out_path))
        finally:
            # Always close the browser, even on failure — an uncaught exception here previously
            # left a headless Chromium process running with its page's polling JS still active,
            # which (via a still-in-flight background case job that page had started) kept
            # occupying Ollama's single concurrency slot and silently starved a later test run.
            browser.close()

    doc = Document(out_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    print("--- exported .docx paragraphs ---")
    print(text)

    missing = [label for label in REQUIRED_LABELS if label not in text]
    table_ok = bool(doc.tables) and [c.text for c in doc.tables[0].rows[0].cells] == [
        "DATE", "PAGE", "RECORD", "SOURCE", "DESCRIPTION",
    ]
    normal = doc.styles["Normal"]
    font_ok = normal.font.name == "Times New Roman" and normal.font.size is not None and normal.font.size.pt == 12

    print("--- check ---")
    print("table headers correct:", table_ok)
    print(f"body font: {normal.font.name} {normal.font.size} (expect Times New Roman, 12pt):", font_ok)
    ok = not missing and table_ok and font_ok
    if missing:
        print(f"FAIL: missing required template labels: {missing}")
    if not ok:
        return 1
    print("PASS: all required template sections/labels present, table structure correct, Times New Roman 12pt")
    return 0


if __name__ == "__main__":
    sys.exit(main())
