#!/usr/bin/env python3
"""Confirms clicking a finding opens an inline editor immediately (see the click handler in
makeFindingEl — clicking both selects AND opens editing in one step), Enter saves it, and the edit
persists (shown as "(edited)" and surfacing in the exported chronology)."""
import sys
from pathlib import Path

from docx import Document
from playwright.sync_api import sync_playwright

BASE = Path(__file__).resolve().parent.parent
ZIP_PATH = BASE / "sample_data" / "case_000_pdfs.zip"
APP_URL = "http://127.0.0.1:5051"


def main():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(f"{APP_URL}/review")
        page.set_input_files("#zipfile", str(ZIP_PATH))
        page.click("#submit-btn")
        page.wait_for_selector("#results", state="visible", timeout=120_000)
        page.wait_for_selector(".finding", timeout=10_000)

        first = page.locator(".finding").first
        original_text = first.locator(".finding-text").inner_text()
        print(f"original text: {original_text!r}")

        # A single click both selects the finding AND opens it for editing immediately (no
        # separate "press Space to edit" step for the mouse — see the comment on the row's click
        # handler in makeFindingEl).
        first.click()
        page.wait_for_timeout(300)

        textarea = first.locator("textarea.finding-edit-textarea")
        if textarea.count() == 0:
            print("FAIL: clicking the finding did not open an inline editor")
            return 1
        textarea_initial_value = textarea.input_value()
        print(f"textarea pre-filled with original text: {textarea_initial_value == original_text}")

        new_text = "EDITED: this fact was corrected by the reviewer."
        textarea.fill(new_text)
        page.keyboard.press("Enter")
        page.wait_for_timeout(300)

        updated_text = first.locator(".finding-text").inner_text()
        edited_tag_visible = first.locator(".finding-edited-tag").is_visible()
        print(f"updated text: {updated_text!r}")
        print(f"'(edited)' tag visible: {edited_tag_visible}")

        # approve it and export, confirm the EDITED text (not the original) shows up
        page.locator("#approve-btn").click()
        page.wait_for_timeout(200)

        with page.expect_download() as dl_info:
            page.locator("#export-btn").click()
            # This fixture can have more than one finding — if only the one just edited got
            # approved, the rest are still unreviewed, which now raises a styled in-app modal
            # (not a native confirm() dialog, see #export-confirm-modal) asking to export anyway.
            page.wait_for_timeout(300)
            proceed_btn = page.locator("#export-confirm-proceed")
            if proceed_btn.is_visible():
                proceed_btn.click()
        download = dl_info.value
        export_path = Path("/tmp/lfa_export_test.docx")
        download.save_as(str(export_path))
        doc = Document(export_path)
        # The finding's text lives in the DATE/PAGE/RECORD/SOURCE/DESCRIPTION table's DESCRIPTION
        # cell, not a plain paragraph — see /export/docx in app.py.
        export_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    export_text += "\n" + cell.text

        ok = (
            textarea_initial_value == original_text
            and updated_text == new_text
            and edited_tag_visible
            and new_text in export_text
            and original_text not in export_text
        )
        print("PASS" if ok else "FAIL")
        browser.close()
        return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
