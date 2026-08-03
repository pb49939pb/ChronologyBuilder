"""
LawFirmAgent — local chronology builder web app.

Upload a .zip of PDFs, they get extracted and sent to a local Ollama model
to build a cited chronology. Everything runs on 127.0.0.1 — nothing leaves
this machine. Prototype only: do not use with real case files/PHI until the
firm approval steps in PRODUCT_DEFINITION.md §7/§9 are done.

The model is asked to return structured JSON with a verbatim supporting quote
per finding, so the browser-side document viewer can locate and highlight the
matching passage in the original PDF. Because of that, uploaded PDFs are now
kept on local disk (system temp dir, under a per-analysis session id) until
the *next* analysis replaces them — a change from the earlier "delete
immediately after the request" behavior. See webapp/README.md.
"""
from __future__ import annotations  # lets "X | None" annotations work on Python 3.9, not just 3.10+

import copy
import hashlib
import io
import json
import os
import re
import shutil
import sys
import tempfile
import threading
import time
import uuid
import zipfile
from datetime import datetime
from pathlib import Path

import numpy as np
import pdfplumber
import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps
from flask import (
    Flask,
    Response,
    g,
    jsonify,
    redirect,
    render_template,
    request,
    send_from_directory,
    stream_with_context,
    url_for,
)

import applog  # local module — structured, PII-safe activity/error logging, see applog.py
import db  # local module — durable SQLite storage, see db.py
import license as licensing  # local module — offline license verification, see license.py
                              # (named `licensing` here only to avoid shadowing the `license`
                              # stdlib-adjacent builtin some tools assume exists; the file itself
                              # is still webapp/license.py)
import pairing  # local module — trust-on-first-use client pairing for remote/tower mode, see pairing.py

def _resource_path(*relative_parts: str) -> Path:
    """Resolves a bundled resource's real path correctly in both dev mode (running from source —
    resource lives in its normal repo-relative place) and a PyInstaller-frozen executable (resource
    lives under sys._MEIPASS's extracted temp directory instead) — see scripts/build_backend.py for
    how prompts/templates/static/easyocr models are added as PyInstaller "datas" at matching paths.
    Frozen detection is PyInstaller's own documented idiom (`sys.frozen` + `sys._MEIPASS`)."""
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS).joinpath(*relative_parts)  # type: ignore[attr-defined]
    return Path(__file__).resolve().parent.parent.joinpath(*relative_parts)


_FROZEN = getattr(sys, "frozen", False)

app = Flask(
    __name__,
    # Flask's default (relative to the app module's own __file__) isn't reliable once frozen —
    # PyInstaller's runtime layout doesn't match the source tree's, so this is pointed explicitly
    # at the bundled copies in that case, and left as Flask's normal default in dev mode.
    template_folder=str(_resource_path("templates")) if _FROZEN else "templates",
    static_folder=str(_resource_path("static")) if _FROZEN else "static",
)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200MB upload cap, sane default for a local tool
app.config["TEMPLATES_AUTO_RELOAD"] = True  # Flask doesn't reload templates by default with debug=False

# 127.0.0.1 is correct for native/bare-metal use (the default) — nothing off this machine can reach
# the app. Set to a real LAN address for remote/tower mode (see TOWER_SETUP.md) — this same value
# also gates _enforce_pairing_key below, since that's exactly the condition under which the app is
# reachable by anything other than itself. Read once at module load rather than only inside
# `__main__` so both that check and the eventual app.run() call agree on the same value.
_BIND_HOST = os.environ.get("LAWFIRMAGENT_BIND_HOST", "127.0.0.1")

# Routes reachable even without a valid license — the license page itself (so there's a way IN),
# and static assets (so that page can actually render with CSS/JS rather than a blank page).
_LICENSE_EXEMPT_ENDPOINTS = {"license_page", "static"}
# Full-page GET routes (render_template, a real browser navigation) — these redirect to /license.
# Every other endpoint (JS-polled JSON APIs, PDF serving, all POST routes) gets a plain 402 JSON
# error instead: redirecting those would hand back an HTML page where the frontend's `resp.json()`
# expects JSON, turning a clear "not licensed" message into a confusing parse error.
_LICENSE_REDIRECT_ENDPOINTS = {"dashboard", "index", "architecture", "case_page"}


@app.before_request
def _enforce_pairing_key():
    # A no-op for every dev/test/local-mode run (all bind loopback) — only activates once this
    # instance is actually reachable off this machine, which is exactly the condition that creates
    # the exposure this exists to close. See pairing.py for what guarantee this does and doesn't
    # provide. Checked before licensing: whether a request is even allowed to reach this box at all
    # is logically prior to whether this box is licensed.
    if _BIND_HOST in ("127.0.0.1", "localhost"):
        return None
    if not pairing.check(request.headers.get("X-Chronology-Pairing-Key")):
        return jsonify({"error": "This device isn't paired with this server."}), 403


@app.before_request
def _enforce_license():
    if request.endpoint in _LICENSE_EXEMPT_ENDPOINTS or request.endpoint is None:
        return None
    status = licensing.get_current_license()
    if status["valid"]:
        return None
    if request.endpoint in _LICENSE_REDIRECT_ENDPOINTS:
        return redirect(url_for("license_page"))
    return jsonify({"error": f"No valid license installed: {status['reason']}"}), 402


@app.context_processor
def _inject_license_display():
    # Makes `license_display` available in every template automatically (see
    # templates/_license_badge.html) without every render_template() call needing to pass it —
    # this is purely a display concern, the actual enforcement is _enforce_license() above.
    return {"license_display": licensing.get_license_display_info()}


# --- Security headers --------------------------------------------------------------------------
# This app has no remote content or third-party scripts of any kind (confirmed: no external URLs
# anywhere in templates/static/*.js/*.css) — every script/style/font/image it ever loads is its own
# bundled static file. That makes a strict CSP nearly free here, unlike a typical web app that has
# to carve out exceptions for a CDN or analytics script. script-src has no 'unsafe-inline'/
# 'unsafe-eval' — confirmed neither pdf.js nor mermaid (the only two non-trivial bundled JS
# libraries) use eval()/new Function() anywhere, and the one inline script this app used to have
# (the update-badge handler) was moved to static/version-badge.js specifically so this could hold.
# style-src keeps 'unsafe-inline' — several templates use inline style="" attributes for simple
# show/hide toggling, and rewriting all of those into classes buys little given style injection
# can't achieve code execution the way script injection can.
@app.after_request
def _set_security_headers(response):
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; script-src 'self'; style-src 'self' 'unsafe-inline'; "
        "img-src 'self' data:; font-src 'self'; connect-src 'self'; worker-src 'self'; "
        "object-src 'none'; base-uri 'none'; form-action 'self'; frame-ancestors 'none'"
    )
    response.headers["X-Content-Type-Options"] = "nosniff"
    return response


# --- Activity/error logging (see applog.py) --------------------------------------------------
# Deliberately logs only method/path/status/duration — never request bodies/query strings, which
# is where user-entered content (plaintiff names, priority hints, etc.) could end up. Every route
# in this app keys state by opaque IDs in the path itself, never free-text in the URL, so the path
# alone is already safe to log in full.

@app.before_request
def _log_request_start():
    g._applog_start = time.time()


@app.after_request
def _log_request_end(response):
    duration_ms = round((time.time() - getattr(g, "_applog_start", time.time())) * 1000)
    applog.log_event(
        "http_request",
        f"{request.method} {request.path} -> {response.status_code}",
        context={
            "method": request.method, "path": request.path,
            "status": response.status_code, "duration_ms": duration_ms,
        },
    )
    return response


@app.errorhandler(Exception)
def _log_unhandled_exception(e):
    # Today, an unhandled exception in a request has no durable record beyond whatever Flask's
    # dev server happens to print to a terminal no one's watching. This is the backstop for every
    # request path, not just the ones with their own explicit try/except (see the per-group/
    # per-file handlers deeper in this file, which log their own tracebacks at the point of
    # failure since they need to keep going rather than fail the whole request).
    applog.log_event(
        "unhandled_exception", f"Unhandled {type(e).__name__} on {request.method} {request.path}",
        level="ERROR", context={"method": request.method, "path": request.path}, exc_info=sys.exc_info(),
    )
    return jsonify({"error": "An unexpected error occurred. Check the logs for details."}), 500


# Single source of truth is the root VERSION file (see scripts/bump_version.py) — read once at
# startup, not per-request, since it only ever changes via a new build/release, never at runtime.
APP_VERSION = _resource_path("VERSION").read_text().strip()


@app.context_processor
def _inject_app_version():
    return {"app_version": APP_VERSION}


OLLAMA_URL = os.environ.get("LAWFIRMAGENT_OLLAMA_URL", "http://127.0.0.1:11434/api/generate")
MODEL = os.environ.get("LAWFIRMAGENT_MODEL", "llama3.1:8b")
# Was 12288 — lowered after a real production timeout (see TEST_RESULTS.md, 2026-08-03): a chunk
# built from this budget took long enough on a real user's Windows laptop to blow past even a
# 30-minute per-call timeout, on multiple chunks in the same run, not a one-off. This value was
# never actually sized for that hardware in the first place — the original 12288 was picked
# conservatively for THIS dev machine's 16GB RAM (see the 2026-07-22 context-window finding this
# same file documents), not verified against anything slower.
#
# IMPORTANT — cannot be lowered arbitrarily: the prompt's own instructional text has grown to
# ~26,000 characters (all the accumulated extraction rules added across many sessions since that
# 12288 default was chosen against a ~10,626-character prompt) — a NAIVE lower value (8192 was
# tried first) actually made MAX_INPUT_CHARS go NEGATIVE, i.e. the fixed instructions alone would
# have exceeded the whole context window before a single character of document text, causing
# silent context overflow (part of the instructions themselves getting truncated) rather than
# anything as obvious as an error. 11264 was chosen specifically to leave a real, non-thin margin
# above that fixed cost (~6,300 chars/~3 pages of actual chunk budget, computed and verified
# directly against the CURRENT prompt, not assumed) — see the hard assertion right after
# MAX_INPUT_CHARS below, added specifically so a future prompt-overhead increase (or a careless
# LAWFIRMAGENT_NUM_CTX override) fails loudly at startup instead of silently repeating this same
# near-miss. Chunking (_pack_into_chunks) already exists specifically so an arbitrarily large CASE
# never depends on a large context window — this only changes how much text goes into any ONE
# call, trading more (smaller, faster) chunks for lower per-call timeout risk.
NUM_CTX = int(os.environ.get("LAWFIRMAGENT_NUM_CTX", "11264"))
CHARS_PER_TOKEN_ESTIMATE = 3.5
OUTPUT_TOKEN_RESERVE = 2000  # structured JSON + per-finding quotes runs longer than plain prose

PROMPT_PATH = _resource_path("prompts", "chronology_prompt_structured.txt")
PROMPT_TEMPLATE = PROMPT_PATH.read_text()

COMPLAINT_PROMPT_PATH = _resource_path("prompts", "complaint_extraction_prompt.txt")
COMPLAINT_PROMPT_TEMPLATE = COMPLAINT_PROMPT_PATH.read_text()

CASE_SYNTHESIS_PROMPT_PATH = _resource_path("prompts", "case_synthesis_prompt.txt")
CASE_SYNTHESIS_PROMPT_TEMPLATE = CASE_SYNTHESIS_PROMPT_PATH.read_text()

DETAIL_LEVELS = {
    "brief": (
        "Report only the most significant findings: major procedures, diagnoses, hospital "
        "admissions/discharges, and only the most clinically or legally significant potential "
        "issues (e.g., a clear delay in escalation, a major documentation gap). Omit routine or "
        "minor events, administrative notes, and low-significance observations. Aim for the "
        "shortest chronology that still captures everything a busy attorney would need to know."
    ),
    "standard": (
        "Report all clinically meaningful events and reasonably significant potential issues. "
        "This is the default level of detail — balance completeness against readability."
    ),
    "detailed": (
        "Report every event mentioned in the documents, however minor — including routine "
        "follow-ups, administrative notes, and any possible concern no matter how small. When in "
        "doubt about whether something is worth including, include it. This level prioritizes "
        "completeness over brevity."
    ),
}
DEFAULT_DETAIL_LEVEL = "standard"

# Reserve space for the prompt's own instructional text (everything except {records}) plus room
# for the model's output, then whatever's left is the budget for document text. Uses the longest
# detail-level instruction so the budget is safe regardless of which one is actually selected.
_longest_detail_instructions = max(DETAIL_LEVELS.values(), key=len)
_prompt_overhead_chars = len(
    PROMPT_TEMPLATE.format(
        records="", detail_instructions=_longest_detail_instructions, case_context=""
    )
)
MAX_INPUT_CHARS = int(
    NUM_CTX * CHARS_PER_TOKEN_ESTIMATE
    - _prompt_overhead_chars
    - OUTPUT_TOKEN_RESERVE * CHARS_PER_TOKEN_ESTIMATE
)

# A real near-miss, not a hypothetical: an earlier attempt to lower NUM_CTX's default computed a
# NEGATIVE MAX_INPUT_CHARS (the prompt's own instructional text alone already exceeded the whole
# context window) — which would have silently produced context overflow / degenerate chunking
# rather than any obvious error. Fails loudly at startup instead, so a future prompt-overhead
# increase (this file's instructional text has only ever grown across sessions) or a careless
# LAWFIRMAGENT_NUM_CTX override surfaces immediately as a clear message, not a silent production
# failure discovered later via a confused user report. 2000 chars is an arbitrary but deliberately
# generous floor — enough for a genuinely tiny real document, not just "technically positive."
if MAX_INPUT_CHARS < 2000:
    raise RuntimeError(
        f"MAX_INPUT_CHARS computed to {MAX_INPUT_CHARS}, too small/negative to safely chunk any "
        f"document — NUM_CTX ({NUM_CTX} tokens) is too small relative to the prompt template's own "
        f"instructional overhead ({_prompt_overhead_chars} characters) plus the output reserve "
        f"({OUTPUT_TOKEN_RESERVE} tokens). Raise LAWFIRMAGENT_NUM_CTX, or shorten the prompt "
        f"template, before starting the app."
    )

# Shared shape for every "dated entry" list (timeline/medications/procedures/diagnoses/labs) — all
# five need the same citation-grounding fields PLUS record_type/author/ordering-reading-provider
# (Olivia's real workflow: title each entry by record type + authoring provider, both bold) and
# at_issue (every mention of a named defendant is "at-issue" and must be visually prominent, even a
# minor one — see the case_context block in the prompt). Kept as required strings with "not
# stated" as the documented escape value, matching this schema's existing omit-rather-than-guess
# convention, rather than introducing optional properties this app hasn't used anywhere else.
#
# Deliberately does NOT also require separate "ordering_provider"/"reading_provider" fields (an
# earlier version did) — verified empirically that going from 4 to 9 required fields per entry made
# the model's output genuinely bimodal on small documents: the exact same call against the same
# short document sometimes produced a full correct extraction and sometimes an entirely empty (but
# schema-valid) result, back-to-back, non-deterministically. Folding the ordering/reading-MD detail
# into "text"/"author" instead of two more always-required fields is the fix attempt — see
# TEST_RESULTS.md for whether this actually resolved it.
def _dated_entry_schema() -> dict:
    return {
        "type": "array",
        "items": {
            "type": "object",
            "properties": {
                "date": {"type": "string"},
                "text": {"type": "string"},
                # A small integer — the number shown in that document's own "### Document N —
                # filename" header — rather than a freeform filename string the model has to
                # reproduce byte-for-byte. Converted back to a real "source_file" filename string
                # immediately after the model call returns (see _resolve_file_ids), before any
                # downstream code (Bates resolution, dedup, the frontend, exports) ever sees it —
                # everything past that conversion point is unchanged and still keys off source_file.
                "file_id": {"type": "integer"},
                "quote": {"type": "string"},
                "record_type": {"type": "string"},
                "author": {"type": "string"},
                "at_issue": {"type": "boolean"},
            },
            "required": ["date", "text", "file_id", "quote", "record_type", "author", "at_issue"],
        },
    }


def _demographics_schema() -> dict:
    """Shared by RESPONSE_SCHEMA (per-chunk extraction) and CASE_SYNTHESIS_SCHEMA (the case-wide
    fill-the-gaps pass) — same seven fields either way."""
    return {
        "type": "object",
        "properties": {
            "dob": {"type": "string"},
            "address": {"type": "string"},
            "social_history": {"type": "string"},
            "past_medical_history": {"type": "string"},
            "family_history": {"type": "string"},
            "surgical_history": {"type": "string"},
            "pcp": {"type": "string"},
        },
        "required": [
            "dob", "address", "social_history", "past_medical_history",
            "family_history", "surgical_history", "pcp",
        ],
    }


RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "timeline": _dated_entry_schema(),
        # A safety-net list of every distinct office visit/patient encounter, same full-narrative
        # richness as timeline entries — added after a real-world case showed the model reliably
        # extracting labs/radiology (simple, explicit, single-fact) while systematically missing
        # whole visit-summary narratives (multi-part synthesis: chief complaint, exam, plan, etc.),
        # the same "structured fact vs. synthesized narrative" gap the other Key Facts lists below
        # already exist to catch for simpler facts. Powers its own "Visits" Key Facts tab so a
        # reviewer can always find/add a visit that didn't make it into the main timeline.
        "visits": _dated_entry_schema(),
        # Atomic, deduplicated facts pulled out separately from the narrative timeline — power the
        # "Key Facts" tabs in the UI, which let the reviewer quickly scan and one-click-add any of
        # these into the reviewed output even if the narrative timeline/summary didn't surface it.
        # Placed right after "timeline" (not last) deliberately: these are simple enumeration tasks
        # that benefit from coming while the model still has its full output budget/attention
        # available, rather than competing with the more interpretive potential_issues/
        # discrepancies/summary sections for whatever's left — see TEST_RESULTS.md, a real test on
        # a document listing 6 medications initially only extracted the first 2 when these fields
        # came last in the schema.
        "medications": _dated_entry_schema(),
        "procedures": _dated_entry_schema(),
        "diagnoses": _dated_entry_schema(),
        "labs": _dated_entry_schema(),
        "potential_issues": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "text": {"type": "string"},
                    "file_id": {"type": "integer"},
                    "quote": {"type": "string"},
                },
                "required": ["text", "file_id", "quote"],
            },
        },
        "discrepancies": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "text": {"type": "string"},
                    "file_ids": {"type": "array", "items": {"type": "integer"}},
                    "quotes": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["text", "file_ids", "quotes"],
            },
        },
        # Case-wide (not per-entry) facts about the patient — stated once, not repeated per finding.
        # See patient_demographics merge handling in _merge_chunk_findings: first non-empty value
        # per subfield wins across chunks (a documented simplification), plus a later case-wide
        # synthesis pass (_synthesize_case_summary) that fills in whatever's still missing after
        # that merge, using visibility across every visit at once instead of one document at a time.
        "patient_demographics": _demographics_schema(),
        # One short, content-derived abbreviation per source document (e.g. "Fairview Op" for an
        # operative report from Fairview) — used in the export/UI's RECORD SOURCE column instead of
        # the raw filename, since a real production's filenames aren't always meaningful. Uniqueness
        # is NOT trusted to the model — enforced deterministically server-side (see
        # _resolve_record_sources), the same reasoning as everywhere else a hard correctness
        # requirement exists in this app. "facility"/"description"/"date" feed the Abbreviations/
        # Record Sources legend line (see _format_record_source_line) — all three degrade gracefully
        # when a document doesn't clearly state them, only "file_id"/"label" are required.
        "record_sources": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "file_id": {"type": "integer"},
                    "label": {"type": "string"},
                    "facility": {"type": "string"},
                    "description": {"type": "string"},
                    "date": {"type": "string"},
                },
                "required": ["file_id", "label"],
            },
        },
        "summary": {"type": "string"},
    },
    "required": [
        "timeline", "visits", "medications", "procedures", "diagnoses", "labs",
        "potential_issues", "discrepancies", "patient_demographics", "record_sources", "summary",
    ],
}

# Matches the "### Document N — filename" header _generate_chronology/the /analyze route already
# write into a chunk/prompt's own assembled text — group 1 is the document's file_id.
_DOC_HEADER_ID_RE = re.compile(r"^### Document (\d+)", re.MULTILINE)


def _valid_file_ids_in_text(text: str) -> list[int]:
    """The distinct file_ids actually present in one call's own already-assembled prompt text,
    read straight back from the "### Document N —" headers rather than tracked separately — so
    this can never silently drift out of sync with what chunk-packing/truncation actually produced
    (an oversized document split into "(part K of M)" sub-blocks keeps the SAME header/number, so
    it's still correctly included exactly once here). Feeds _response_schema_for_ids: the schema's
    enum for this one call is restricted to exactly these ids, so the model cannot reference a
    document it wasn't actually given in this call, or one packed into a different chunk."""
    return sorted({int(m) for m in _DOC_HEADER_ID_RE.findall(text)})


def _response_schema_for_ids(valid_file_ids: list[int]) -> dict:
    """A deep copy of RESPONSE_SCHEMA with every file_id/file_ids field restricted via JSON-schema
    "enum" to exactly the document ids present in one call's own prompt text (see
    _valid_file_ids_in_text). Confirmed via a standalone smoke test that Ollama's grammar-
    constrained decoding genuinely enforces "enum" on an integer field (not just "type"/
    "required", the only constraints this schema exercised before) — even a prompt actively trying
    to induce an out-of-enum value could not produce one. Built fresh per call rather than mutating
    RESPONSE_SCHEMA in place, since that module-level constant must stay a stable, reusable
    template — deep-copying a schema this small is cheap."""
    schema = copy.deepcopy(RESPONSE_SCHEMA)
    enum_int = {"type": "integer", "enum": valid_file_ids}
    for field in _DATED_LIST_FIELDS + ("potential_issues",):
        schema["properties"][field]["items"]["properties"]["file_id"] = enum_int
    schema["properties"]["discrepancies"]["items"]["properties"]["file_ids"] = {
        "type": "array", "items": enum_int,
    }
    schema["properties"]["record_sources"]["items"]["properties"]["file_id"] = enum_int
    return schema


COMPLAINT_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "plaintiff_name": {"type": "string"},
        "dol": {"type": "string"},
        "facts_summary": {"type": "string"},
        "quote": {"type": "string"},
    },
    # Deliberately no "defendant_names" — the reviewer enters these herself when starting the case
    # (see the case-start form), rather than trusting the model to identify them from the complaint.
    "required": ["plaintiff_name", "dol", "facts_summary", "quote"],
}

CASE_SYNTHESIS_SCHEMA = {
    "type": "object",
    "properties": {
        "dol": {"type": "string"},
        "summary": {"type": "string"},
        "patient_demographics": _demographics_schema(),
    },
    "required": ["dol", "summary", "patient_demographics"],
}

PORT = int(os.environ.get("LAWFIRMAGENT_PORT", "5050"))

# Only the most recent analysis's PDFs are kept on disk — starting a new one deletes the last.
# Suffixed by port so a separate test-server instance (see testing/) never collides with or wipes
# out whatever session is active on the normal port a person is actually using.
SESSIONS_DIR = Path(tempfile.gettempdir()) / f"lawfirmagent_sessions_{PORT}"


def _clear_old_sessions():
    shutil.rmtree(SESSIONS_DIR, ignore_errors=True)
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)


# Case Mode jobs (see /case routes) get durable, long-term storage — deliberately NOT under
# tempfile.gettempdir() (that's only for SESSIONS_DIR above, the plain single-upload /review flow,
# which has no job_id and was never meant to be looked up again later). A case's manifest lives in
# the SQLite database (see db.py) at a per-OS application-data directory that survives OS temp
# cleanup, port changes, and app restarts — "pull up a chronology from a month ago" requires this.
# Source PDFs (for the document viewer/highlighting) get the same durability treatment: copied into
# a directory under that same application-data location, not a port-namespaced temp dir, so a case's
# highlighting keeps working indefinitely, not just until the next OS temp sweep.
BATCH_SESSIONS_DIR = db.get_app_data_dir() / "case_sessions"
BATCH_SESSIONS_DIR.mkdir(parents=True, exist_ok=True)

# Where an uploaded case folder's original files land — see case_start()/case_rescan() below. This
# app used to have Flask read case folders directly off whatever disk it happened to be running on
# (a path typed/browsed to in the UI), which required Flask's OS user to already have broad
# filesystem read permission to arbitrary folders (macOS Full Disk Access, or the Windows-service
# equivalent) — fragile, and a fundamentally different trust model than the single-zip-upload flow,
# which just works because the BROWSER's own native picker is what's granting file access, not the
# server. Case folders are now uploaded the same way: a native <input type="file" webkitdirectory>
# picker in case.html/index.html, multipart-uploaded to Flask, which needs no special permissions of
# its own to read them back from here.
CASE_SOURCES_DIR = db.get_app_data_dir() / "case_sources"
CASE_SOURCES_DIR.mkdir(parents=True, exist_ok=True)

db.init_db()
applog.cleanup_old_logs()  # best-effort, bounded to 90 days — see applog.py

_batch_jobs_lock = threading.Lock()
_batch_jobs: dict = {}  # job_id -> manifest dict (the SAME mutable object _run_case_job mutates in
                         # place across its steps) — kept as an in-memory cache/source-of-truth for
                         # a currently-running job so a concurrent read sees fully current state even
                         # between two _save_batch_manifest flushes to disk; NOT just a cache of what's
                         # in SQLite, don't remove this or a running job's progress reads go stale.


def _save_batch_manifest(job_id: str, manifest: dict) -> None:
    with _batch_jobs_lock:
        _batch_jobs[job_id] = manifest
    db.save_case_manifest(job_id, manifest)


def _load_batch_manifest(job_id: str) -> dict | None:
    return db.load_case_manifest(job_id)


def _get_batch_manifest(job_id: str) -> dict | None:
    with _batch_jobs_lock:
        manifest = _batch_jobs.get(job_id)
    return manifest if manifest is not None else _load_batch_manifest(job_id)


# --- Patient identification, for grouping a folder's mixed documents by patient ---
#
# Filename-based guess is tried first (cheap, no need to read the file) and falls back to
# content-based extraction — real EHR exports we've seen (e.g. "Visit Summary for
# 06:25:2026.pdf") don't embed patient names in filenames at all, so in practice the content-based
# path does most of the work; the filename check stays as a fast path for folders that do name
# files after patients.

_FILENAME_NAME_PATTERN = re.compile(r"^\s*([A-Za-z'\-]+)\s*,\s*([A-Za-z'\-]+)")

_PATIENT_NAME_PATTERNS = [
    # "Last Name: Susalla    First Name: Olivia" — same pattern our own OCR/text pipeline has
    # actually encountered (see the "Visit Summary" documents referenced in TEST_RESULTS.md).
    re.compile(
        r"Last\s*Name\s*:\s*([A-Za-z'\-]+).{0,80}?First\s*Name\s*:\s*([A-Za-z'\-]+)",
        re.IGNORECASE | re.DOTALL,
    ),
    re.compile(r"Patient\s*Name\s*:\s*([A-Za-z'\-]+)\s+([A-Za-z'\-]+)", re.IGNORECASE),
    # Generic "Name: First Last" — only reached if the two more specific patterns above didn't
    # match, so this won't misfire on "Patient Name:"/"Last Name:"/"First Name:" lines.
    re.compile(r"(?:^|\n)\s*Name\s*:\s*([A-Za-z'\-]+)\s+([A-Za-z'\-]+)", re.IGNORECASE),
]

_DOB_PATTERN = re.compile(r"DOB\s*:?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})", re.IGNORECASE)


def _guess_patient_from_filename(filename: str) -> str | None:
    m = _FILENAME_NAME_PATTERN.match(Path(filename).stem)
    if m:
        return f"{m.group(2)} {m.group(1)}"
    return None


def _extract_patient_identity(text: str) -> dict:
    """Best-effort regex extraction of a patient's name and DOB (when present) from a document's
    extracted text. DOB is used as a secondary grouping key so two different patients who happen
    to share a name don't get merged into one chronology — a real, if less common, failure mode
    worth guarding against explicitly, not just an edge case to ignore."""
    name = None
    for i, pattern in enumerate(_PATIENT_NAME_PATTERNS):
        m = pattern.search(text)
        if m:
            first, second = m.groups()
            # The first pattern captures (last, first); the other two capture (first, last).
            name = f"{second} {first}" if i == 0 else f"{first} {second}"
            break
    dob_match = _DOB_PATTERN.search(text)
    return {"name": name, "dob": dob_match.group(1) if dob_match else None}


def _normalize_patient_key(name: str | None, dob: str | None, fallback_id: str = "") -> str:
    """`fallback_id` (typically a hash of the filename) MUST be unique per unidentifiable document
    — found necessary against a real batch of professional sample chronologies where every one of
    7 completely unrelated documents used a literal placeholder header ("Patient Name DOB:
    MM/DD/YYYY", no real name filled in). Regex correctly declined to extract that as a name
    (no colon follows "Patient Name" in that malformed header), but a single shared "unidentified"
    key for every such case would have merged all 7 unrelated documents into one 364-page blob —
    actively worse than just admitting each one couldn't be identified. Never share this key across
    genuinely different, unidentifiable documents.
    """
    if not name:
        return f"unidentified_{fallback_id}" if fallback_id else "unidentified"
    norm_name = re.sub(r"[^a-z]", "", name.lower())
    if dob:
        norm_dob = re.sub(r"[^0-9]", "", dob)
        return f"{norm_name}_{norm_dob}"
    return norm_name


def _same_person(name_a: str | None, name_b: str | None) -> bool:
    """Loose match used to decide whether a document found in a case folder is about the case's own
    plaintiff (stays in the primary chronology) or a different, incidentally-referenced individual
    (gets its own secondary group — see _run_case_job). Substring containment (not exact equality)
    on purpose: real documents write a name as "Smith, Jane" vs "Jane M. Smith" vs just "Jane Smith"
    often enough that exact matching would wrongly split one person's own records into two groups.
    """
    norm = lambda n: re.sub(r"[^a-z]", "", (n or "").lower())
    a, b = norm(name_a), norm(name_b)
    return bool(a) and bool(b) and (a == b or a in b or b in a)


# --- Complaint / Notice of Intent (NOI) detection — the document a case is built around ---
#
# A case's folder normally has a "pleadings" subfolder containing the Complaint or NOI that states
# the factual basis of the suit (including the Date of Loss) and names the defendant(s). Detected by
# filename pattern, preferring that subfolder, so the rest of the case-analysis prompt can be told
# who the defendant is and what to look for before it ever reads a medical record — per the user's
# explicit instruction to always start a case with the complaint file.
_COMPLAINT_FILENAME_PATTERN = re.compile(
    r"complaint|notice[\s_-]*of[\s_-]*intent|\bnoi\b", re.IGNORECASE
)


def _find_complaint_file(pdf_paths: list, folder: Path) -> Path | None:
    """Best-effort: prefer a match inside a 'pleadings' subfolder (case-insensitive dir name), else
    anywhere in the tree. Returns None (never raises) if nothing matches — the caller proceeds
    without case context and surfaces a clear warning rather than blocking the case."""
    candidates = [p for p in pdf_paths if _COMPLAINT_FILENAME_PATTERN.search(p.stem)]
    if not candidates:
        return None
    in_pleadings = [
        p for p in candidates
        if any(part.lower() == "pleadings" for part in p.relative_to(folder).parts[:-1])
    ]
    return (in_pleadings or candidates)[0]


def _extract_complaint_info(pdf_path: Path) -> dict:
    """One schema-constrained Ollama call against the Complaint/NOI's extracted text — same
    verbatim-quote-grounding discipline as the main chronology extraction, just a different, much
    smaller schema (see COMPLAINT_RESPONSE_SCHEMA/complaint_extraction_prompt.txt). Raises on a
    request/JSON error, same as _call_model_for_chunk — the caller decides how to degrade (see
    _run_case_job, which proceeds without case context and warns rather than failing the case)."""
    text, _, _ = pdf_to_text(pdf_path)
    prompt = COMPLAINT_PROMPT_TEMPLATE.format(complaint_text=text.strip())
    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "format": COMPLAINT_RESPONSE_SCHEMA,
        "options": {"num_ctx": NUM_CTX},
    }
    resp = requests.post(OLLAMA_URL, json=payload, timeout=1800)
    resp.raise_for_status()
    return json.loads(resp.json()["response"])


def _synthesize_case_summary(
    timeline: list, visits: list, demographics_so_far: dict,
    complaint_dol: str | None, facts_summary: str | None, case_context: str,
) -> dict:
    """One schema-constrained Ollama call over the PRIMARY group's already-merged, deduped
    timeline/visits — not raw documents — run once chunking/merging is entirely done. This is what
    lets it see the WHOLE case at once (each chunk's own model call, by contrast, only ever sees its
    own slice): determining a real Date of Loss grounded in the actual timeline rather than just the
    complaint's own (often vague/single-date) statement, writing one coherent summary instead of
    _merge_chunk_findings' `[Part 1]/[Part 2]` concatenation (that function's own docstring already
    documents as unable to notice a cross-chunk trend), and filling in demographics fields still
    "not stated" after the per-chunk merge — a real, observed gap: scattered demographic mentions
    (a social history aside in one visit, a family history comment in another) often don't get
    picked up analyzing one document/chunk at a time, but are easier to piece together with every
    visit narrative visible at once. Feasible regardless of case size specifically because the input
    here is already-condensed extracted facts, not raw document text — the same insight that
    motivated chunking in the first place. Raises on a request/JSON error, same as
    _extract_complaint_info — the caller (_run_case_job/_run_case_rescan) decides how to degrade.
    """
    def entries_text(items: list) -> str:
        return "\n".join(
            f"- {item.get('date') or 'not stated'}: {item.get('text') or ''}" for item in items
        ) or "(none)"

    demographics_so_far = demographics_so_far or {}
    demo_lines = "\n".join(
        f"- {field}: {demographics_so_far.get(field) or 'not stated'}"
        for field in ("dob", "address", "social_history", "past_medical_history",
                       "family_history", "surgical_history", "pcp")
    )
    prompt = CASE_SYNTHESIS_PROMPT_TEMPLATE.format(
        case_context=case_context or "(no case context available)",
        complaint_dol=complaint_dol or "not stated",
        facts_summary=facts_summary or "not stated",
        timeline_text=entries_text(timeline),
        visits_text=entries_text(visits),
        demographics_so_far=demo_lines,
    )
    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "format": CASE_SYNTHESIS_SCHEMA,
        "options": {"num_ctx": NUM_CTX},
    }
    resp = requests.post(OLLAMA_URL, json=payload, timeout=1800)
    resp.raise_for_status()
    return json.loads(resp.json()["response"])


def _build_case_context(
    plaintiff_name: str, defendant_names: list, dol: str | None,
    facts_summary: str | None, priority_hint: str,
) -> str:
    lines = [f"This case is being analyzed on behalf of: {plaintiff_name} (Plaintiff)."]
    if defendant_names:
        lines.append(f"Defendant(s) we represent: {', '.join(defendant_names)}.")
    else:
        lines.append(
            "No defendant name is known for this case — always answer false for at_issue rather "
            "than guessing who might be involved."
        )
    if dol and _realish(dol):
        lines.append(f"Date of Loss (DOL): {dol}.")
    if facts_summary and _realish(facts_summary):
        lines.append(f"Factual basis of the claim (from the Complaint/Notice of Intent): {facts_summary}")
    if priority_hint and priority_hint.strip():
        lines.append(
            f"The reviewer flagged these as especially important to focus on: {priority_hint.strip()}. "
            "Give them extra attention, but still fully analyze every document."
        )
    return "\n".join(lines)


def _build_incremental_context(manifest: dict, existing_primary_group: dict) -> str:
    """Prior-context block for reprocessing NEWLY-ADDED files against an already-processed case
    (see _run_case_rescan) — the same case-level facts _build_case_context already provides, plus a
    compact summary of what's already known from the files already processed, so the model can be
    "intelligent about what it's extracting" (the user's own words) instead of starting from a blank
    slate. Deliberately NOT the full prior chronology text — that could be huge on a large case and
    blow the context budget; just enough to avoid redundant re-derivation and give the model
    continuity (e.g. already-known demographics, what sources have already been reviewed)."""
    base = _build_case_context(
        manifest["plaintiff_name"], manifest.get("defendant_names") or [],
        manifest.get("dol"), manifest.get("facts_summary"), manifest.get("priority_hint", ""),
    )
    demo = (existing_primary_group.get("findings") or {}).get("patient_demographics") or {}
    known_demo = "; ".join(f"{k}: {v}" for k, v in demo.items() if _realish(v))
    known_labels = ", ".join((existing_primary_group.get("record_source_labels") or {}).values())
    lines = [base]
    if known_demo:
        lines.append(f"Already known about this patient from previously-processed records: {known_demo}.")
    if known_labels:
        lines.append(
            f"Already-processed record sources (for context only, these are NOT the documents you're "
            f"reading now): {known_labels}."
        )
    lines.append(
        "The documents below are NEWLY ADDED to this case, on top of records already reviewed. "
        "Analyze only what's in the documents actually provided to you now."
    )
    return "\n".join(lines)


def _merge_incremental_group_result(existing_group: dict, new_result: dict) -> dict:
    """Combines a group's already-durable findings with a fresh result covering only newly-added
    files, for _run_case_rescan. Reuses _merge_chunk_findings (already built to concatenate N
    independently-generated findings dicts) treating "existing findings" and "new findings" as two
    chunks — existing MUST be passed first so _assign_finding_ids' collision-disambiguation (see
    that function) never reassigns an already-reviewed item's id, even in the rare case a brand-new
    item's content happens to hash-collide with an old one.

    Deliberately does NOT re-run _resolve_bates on the combined set: each pass already resolved its
    own items' Bates numbers correctly against the exact text it was given (old items against old
    files' text at original-creation time, new items against new files' text just now) — re-running
    it on the union would need every old file's text again, not cheaply available here, and buys
    nothing since nothing about an old item's already-resolved Bates number needs to change.
    """
    combined = _merge_chunk_findings([existing_group.get("findings") or {}, new_result["findings"]])
    _dedupe_all_with_multi_source(combined)
    _assign_finding_ids(combined)  # recompute after merge/dedup — see _assign_finding_ids
    all_source_filenames = (existing_group.get("source_filenames") or []) + new_result["source_filenames"]
    existing_stats = existing_group.get("stats") or {}
    return {
        "findings": combined,
        "source_filenames": all_source_filenames,
        "record_source_labels": _resolve_record_sources(combined, all_source_filenames),
        "page_counts": {**(existing_group.get("page_counts") or {}), **new_result["page_counts"]},
        "warnings": (existing_group.get("warnings") or []) + new_result["warnings"],
        "ocr_files": (existing_group.get("ocr_files") or []) + new_result["ocr_files"],
        "stats": {
            **new_result["stats"],
            "documents_processed": existing_stats.get("documents_processed", 0) + new_result["stats"]["documents_processed"],
            "total_pages": existing_stats.get("total_pages", 0) + new_result["stats"]["total_pages"],
        },
    }


def safe_extract_pdfs(zip_path: Path, extract_dir: Path) -> list[Path]:
    """Extract only .pdf members from the zip, guarding against zip-slip path traversal."""
    extract_dir = extract_dir.resolve()
    pdf_paths = []
    with zipfile.ZipFile(zip_path) as zf:
        for info in zf.infolist():
            name = info.filename
            if info.is_dir() or "__MACOSX" in name or not name.lower().endswith(".pdf"):
                continue
            dest = (extract_dir / Path(name).name).resolve()
            if extract_dir not in dest.parents and dest != extract_dir:
                continue
            with zf.open(info) as src, open(dest, "wb") as out:
                shutil.copyfileobj(src, out)
            pdf_paths.append(dest)
    return pdf_paths


MIN_CHARS_PER_PAGE = 20  # below this, a page's text layer is treated as unreliable/absent

# EasyOCR's Reader loads its detection/recognition models on first use (a few seconds) — build
# it once lazily rather than per-page, and only if a page actually needs the OCR fallback.
_OCR_READER = None


def _get_ocr_reader():
    global _OCR_READER
    if _OCR_READER is None:
        import easyocr

        # gpu=True lets easyocr use Apple Silicon's Metal GPU (via PyTorch's "mps" backend) when
        # available, falling back to CPU automatically otherwise — confirmed via a benchmark
        # across 6 synthetic hard-case test PDFs (testing/ocr_benchmark.py,
        # testing/ocr_results/benchmark_results.json): 3-4x faster than CPU with IDENTICAL
        # accuracy on every test case (same model, same weights, just a different compute
        # backend) — a clean win, not a speed/accuracy tradeoff.
        #
        # In dev mode, easyocr's own default (~/.EasyOCR, downloading on first use if missing) is
        # left untouched. Frozen builds must never touch the network (see PRODUCT_DEFINITION.md
        # §7's no-egress requirement) — model_storage_directory points at the copy bundled by
        # scripts/build_backend.py instead, with downloads disabled outright so a missing/corrupt
        # bundled file fails loudly rather than silently reaching out to the internet.
        if _FROZEN:
            _OCR_READER = easyocr.Reader(
                ["en"], gpu=True, verbose=False,
                model_storage_directory=str(_resource_path("easyocr_models")),
                download_enabled=False,
            )
        else:
            _OCR_READER = easyocr.Reader(["en"], gpu=True, verbose=False)
    return _OCR_READER


def _deskew(image):
    """Cheap deskew: tries a small range of rotation angles and picks the one that maximizes the
    variance of horizontal row-darkness projections — text lines produce sharp peaks/troughs when
    correctly horizontal, and a flatter profile when still skewed. Benchmarked as the single
    biggest OCR accuracy win found (testing/ocr_benchmark.py): a page scanned at a ~4 degree skew
    went from 42% to 98% word accuracy after this step, with negligible cost on already-straight
    pages."""
    gray = ImageOps.grayscale(image)
    best_angle, best_score = 0, -1
    for angle in np.arange(-6, 6.5, 0.5):
        rotated = gray.rotate(angle, fillcolor=255, resample=Image.BICUBIC)
        row_sums = np.array(rotated, dtype=np.float32).sum(axis=1)
        score = np.var(row_sums)
        if score > best_score:
            best_score, best_angle = score, angle
    if best_angle == 0:
        return image
    return image.rotate(best_angle, fillcolor=(255, 255, 255), resample=Image.BICUBIC)


def _ocr_page(page) -> str:
    """Rasterizes a pdfplumber page and runs OCR on it — the fallback for scanned/faxed pages
    that have no usable text layer at all.

    Settings below were chosen from a benchmark across 6 synthetic hard-case test PDFs (clean
    typed, two handwriting-style tiers, faxed/noisy, skewed, and a worst-case combination — see
    testing/ocr_benchmark.py and testing/ocr_results/benchmark_results.json), not assumed:
    200 DPI (300 DPI was consistently no better and sometimes slightly worse, as well as slower —
    counter-intuitive but measured, not guessed), plus deskew + auto-contrast preprocessing (the
    single biggest accuracy win found, and cheap: no ML, pure PIL/numpy).
    """
    image = page.to_image(resolution=200).original.convert("RGB")
    image = _deskew(image)
    image = ImageOps.autocontrast(image, cutoff=1)
    reader = _get_ocr_reader()
    lines = reader.readtext(np.array(image), detail=0, paragraph=True)
    return "\n".join(lines)


# A real column gutter is a wide gap between words; ordinary word spacing is a few points —
# used both to find the split point and to tell a genuine two-column row apart from a full-width
# line that merely has a word starting past the split.
_COLUMN_GUTTER_MIN_PT = 12


def _extract_table_aware_text(page, tables) -> str:
    """Extracts a page that has genuine table structure (real ruled lines or shaded/bordered
    cells — not just visually-aligned text) in correct ROW-by-row order, with any text before/
    between/after the table(s) extracted normally.

    This is the opposite correction from the column-reconstruction logic below, and both are
    needed for different real documents: a genuine table's columns (e.g. a professional
    chronology's "DATE | FACILITY | MEDICAL EVENTS | PDF REF" row) belong TOGETHER per row — a
    date and its own event description must stay paired — whereas two independent prose columns
    on a form (e.g. demographics on the left, vitals on the right) do NOT correspond row-by-row at
    all and need the opposite treatment. Found necessary by testing against real professional
    chronology samples (see TEST_RESULTS.md): running the column-reconstruction logic on a
    genuine table separated every date from its own event description into two disconnected
    blocks, which is a worse, actively harmful misreading — confirmed via `page.find_tables()`
    returning real structural evidence (rects/lines) on these pages, versus none on the
    independent-two-column EHR export that motivated the column-reconstruction fix in the first
    place, so table detection is a reliable, principled way to pick the right treatment per page.
    """
    parts = []
    cursor_top = 0.0
    for table in sorted(tables, key=lambda t: t.bbox[1]):
        x0, top, x1, bottom = table.bbox
        if top > cursor_top + 2:
            above = page.within_bbox((0, cursor_top, page.width, top))
            above_text = (above.extract_text() or "").strip()
            if above_text:
                parts.append(above_text)
        for row in table.extract():
            cells = [(c or "").strip() for c in row]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
        cursor_top = bottom

    if cursor_top < page.height - 2:
        below = page.within_bbox((0, cursor_top, page.width, page.height))
        below_text = (below.extract_text() or "").strip()
        if below_text:
            parts.append(below_text)

    return "\n".join(parts)


def _extract_page_text(page) -> str:
    """Extracts a page's text in proper reading order, rather than trusting pdfplumber's default
    (which sorts primarily by vertical position and badly interleaves content whenever two
    unrelated regions overlap in y-range).

    Checks for genuine table structure first (see _extract_table_aware_text) — real ruled lines or
    shaded/bordered cells, not just visually-aligned columns — since a table's columns need
    row-by-row reading, the opposite of what the column-reconstruction logic below does. Only
    falls through to column-reconstruction (for two independently-flowing prose columns, like a
    form's demographics-left/vitals-right layout) when no real table is present.

    Found necessary on a real multi-column EHR export: a 6-item "Current Medications" list in the
    left column was getting broken up line-by-line by unrelated right-column content (vitals,
    visit notes) that merely happened to sit at the same page height — the model then only picked
    up the first couple of medications before the interleaved unrelated text looked like a topic
    change. Confirmed server-side extraction (this function) was the actual cause, not the model:
    pdfplumber's plain `extract_text()` on that exact page showed the same fragmentation.

    Detects whether the page has a consistent vertical gutter separating two columns of content
    (allowing a few full-width "divider" rows, like a section header, to cross it); if so,
    reconstructs reading order as "all of the left column, top to bottom, then all of the right
    column, top to bottom" rather than row-by-row. Falls back to plain `extract_text()` for
    genuinely single-column pages, or anything too sparse to analyze confidently.
    """
    tables = page.find_tables()
    if tables:
        return _extract_table_aware_text(page, tables)

    words = page.extract_words(x_tolerance=1.5, keep_blank_chars=False)
    if len(words) < 20:
        return page.extract_text() or ""

    rows = {}
    for w in words:
        key = round(w["top"] / 3) * 3  # bucket words into rows by vertical position
        rows.setdefault(key, []).append(w)
    row_keys = sorted(rows.keys())
    if len(row_keys) < 8:
        return page.extract_text() or ""

    # Find the x that separates the page into two columns most cleanly: the candidate where the
    # fewest rows have a word actually straddling it. A genuine two-column page will have one x
    # where almost no row crosses it at all (its gutter); a single-column page won't.
    lo, hi = page.width * 0.25, page.width * 0.75
    best_x, best_clean_frac = None, 0.0
    x = lo
    while x <= hi:
        crossed = sum(
            1 for row_words in rows.values() if any(w["x0"] < x < w["x1"] for w in row_words)
        )
        clean_frac = 1 - (crossed / len(row_keys))
        if clean_frac > best_clean_frac:
            best_clean_frac, best_x = clean_frac, x
        x += 2

    if best_x is None or best_clean_frac < 0.85:
        return page.extract_text() or ""  # no clean split found — treat as single-column

    def is_divider_row(row_words):
        """True only for a row whose content genuinely spans both columns as one continuous
        unit (e.g. a full-width section header) — never true for a row with content on only one
        side, which must always be appended to that side without flushing the other, or the two
        columns' accumulated text gets needlessly interleaved again."""
        left = [w for w in row_words if w["x0"] < best_x]
        right = [w for w in row_words if w["x0"] >= best_x]
        if not left or not right:
            return False
        if any(w["x0"] < best_x < w["x1"] for w in row_words):
            return True  # a word literally straddles the split
        gap = min(w["x0"] for w in right) - max(w["x1"] for w in left)
        return gap < _COLUMN_GUTTER_MIN_PT

    lines = []
    left_block, right_block = [], []

    def flush():
        lines.extend(left_block)
        lines.extend(right_block)
        left_block.clear()
        right_block.clear()

    for key in row_keys:
        row_words = sorted(rows[key], key=lambda w: w["x0"])
        if is_divider_row(row_words):
            flush()
            lines.append(" ".join(w["text"] for w in row_words))
        else:
            left = [w["text"] for w in row_words if w["x0"] < best_x]
            right = [w["text"] for w in row_words if w["x0"] >= best_x]
            if left:
                left_block.append(" ".join(left))
            if right:
                right_block.append(" ".join(right))
    flush()

    return "\n".join(lines)


# Every page of every real document this app processes is guaranteed to carry a Bates stamp
# (confirmed with the firm) — a legal/medical-records production convention: a unique, sequential,
# zero-padded identifier stamped in the bottom margin of each page (left, center, or right), used
# to cite a specific page unambiguously instead of an ambiguous "page N of this PDF" (which page
# numbering scheme? within which of possibly several merged source documents?). See TEST_RESULTS.md
# for the research this was based on. Matches formats like "SMITH000123", "MED_000045", "000001" —
# an optional letter/underscore prefix followed by a zero-padded digit run that always starts with
# a leading zero (per the firm's own numbering convention).
_BATES_PATTERN = re.compile(r"\b[A-Za-z]{0,6}[-_ ]?0\d{3,9}\b")
_BATES_BOTTOM_MARGIN_FRAC = 0.12  # bottom 12% of the page — "always at the bottom", per the firm


def _extract_bates_number(page) -> str | None:
    """Finds the Bates stamp in a page's bottom margin. Position (left/center/right) deliberately
    doesn't matter — only vertical position does — since the firm's own stamps vary in horizontal
    placement. Returns None if nothing matching the expected format is found there (e.g. a scanned
    page OCR didn't pick it up cleanly) rather than guessing; the caller/prompt treats a page with
    no resolvable Bates number as one whose facts can't be cited, the same "omit rather than guess"
    treatment as an unresolvable quote or source file.
    """
    bottom_cutoff = page.height * (1 - _BATES_BOTTOM_MARGIN_FRAC)
    try:
        words = page.extract_words(x_tolerance=1.5, keep_blank_chars=False)
    except Exception:
        return None
    for w in words:
        if w["top"] >= bottom_cutoff and _BATES_PATTERN.fullmatch(w["text"]):
            return w["text"]
    return None


def pdf_to_text(pdf_path: Path) -> tuple[str, int, int]:
    """Extracts text page by page, falling back to OCR for pages whose text layer is missing or
    too sparse to be useful (e.g. scanned/faxed records with no embedded text at all).

    pdfplumber is used instead of pypdf for the primary extraction — it handles multi-column
    layouts and irregular spacing noticeably better, which matters for medical records. Reading
    order within a page is reconstructed column-aware (see _extract_page_text) rather than using
    pdfplumber's raw default, which can badly interleave two independent columns.

    Each page's own Bates number (see _extract_bates_number) is embedded directly in the text
    stream as a `[BATES: XXXXX]` marker at the start of that page's content — the model reads this
    the same way it reads everything else, and is told (see the prompt) to cite the specific Bates
    number a fact came from rather than an ambiguous page index. Feeding it in-band like this,
    rather than only discovering it client-side after a reviewer happens to open a finding's source,
    means every approved finding has a real citeable Bates number by export time, not just the ones
    someone clicked on.

    Returns (combined_text, num_pages, num_ocr_pages) so the caller can report when OCR kicked in.
    """
    pages_text = []
    num_ocr_pages = 0
    with pdfplumber.open(str(pdf_path)) as pdf:
        num_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            text = _extract_page_text(page).strip()
            if len(text) < MIN_CHARS_PER_PAGE:
                try:
                    ocr_text = _ocr_page(page).strip()
                except Exception:
                    ocr_text = ""
                if len(ocr_text) > len(text):
                    text = ocr_text
                    num_ocr_pages += 1
            # [PAGE N] is always present (gives every page a segment boundary + a real page number
            # regardless of whether it's Bates-stamped); [BATES: X] is added ONLY when a real stamp
            # is actually found — no fabricated placeholder. A page with no genuine Bates stamp used
            # to get a literal "UNKNOWN" string embedded here, which then flowed through citation
            # resolution as if it were a real Bates number — a real bug, not hypothetical, found
            # while adding the page-number fallback this replaces it with.
            bates = _extract_bates_number(page)
            marker = f"[PAGE {page_num}]\n" + (f"[BATES: {bates}]\n" if bates else "")
            pages_text.append(f"{marker}{text}")
    return "\n".join(pages_text), num_pages, num_ocr_pages


def _normalize_filename(name: str) -> str:
    return re.sub(r"[^a-z0-9]", "", name.lower())


def _find_closest_pdf(session_dir: Path, requested_name: str) -> Path | None:
    """Best-effort fuzzy match against the PDFs actually in this session.

    The model doesn't always reproduce a complex filename byte-for-byte in its `source_file`
    output — short simple names like "record_01.pdf" come back verbatim, but longer descriptive
    names with colons, em-dashes, or unusual punctuation are more error-prone. Rather than 404 on
    any minor discrepancy, normalize away case/punctuation/whitespace on both sides and match
    against what's actually on disk.
    """
    normalized_requested = _normalize_filename(requested_name)
    if not normalized_requested:
        return None
    candidates = [p for p in session_dir.iterdir() if p.suffix.lower() == ".pdf"]
    for p in candidates:
        if _normalize_filename(p.name) == normalized_requested:
            return p
    for p in candidates:
        normalized_candidate = _normalize_filename(p.name)
        if normalized_candidate and (
            normalized_candidate in normalized_requested or normalized_requested in normalized_candidate
        ):
            return p
    return None


@app.route("/license", methods=["GET", "POST"])
def license_page():
    error = None
    success = None
    if request.method == "POST":
        token = request.form.get("token", "").strip()
        if not token:
            error = "Please paste in a license key."
        else:
            try:
                payload = licensing.install_license(token)
                tier_label = licensing.TIER_LABELS.get(payload["tier"], payload["tier"])
                success = f"Success — license key accepted. Licensed to {payload['customer']}, {tier_label} tier."
                # Tier only, deliberately never the customer/firm name string — see applog.py's
                # "what's safe to log" note.
                applog.log_event("license_installed", "License key accepted", context={"tier": payload["tier"]})
            except licensing.InvalidLicense as e:
                error = str(e)
                applog.log_event("license_rejected", f"License key rejected: {e}", level="WARNING")
    status = licensing.get_current_license()
    return render_template("license.html", status=status, error=error, success=success)


@app.route("/")
def dashboard():
    return render_template("dashboard.html")


@app.route("/review")
def index():
    return render_template("index.html", model=MODEL)


@app.route("/architecture")
def architecture():
    return render_template("architecture.html")


_AT_ISSUE_RGB = RGBColor(150, 40, 30)


def _add_page_number_field(paragraph, instr: str) -> None:
    """Inserts a live Word field code (e.g. "PAGE" or "NUMPAGES") into `paragraph` — python-docx has
    no high-level API for these, since the actual number only exists once Word itself paginates the
    document on open/print. This builds the same raw OOXML field-code sequence Word's own Insert
    Page Number feature produces, so "Page X of Y" stays accurate no matter how long the export is."""
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_separate)
    run._r.append(fld_end)


@app.route("/export/docx", methods=["POST"])
def export_docx():
    """Renders the reviewer's approved chronology as a real Microsoft Word document (Times New
    Roman, 12pt body text throughout — the firm's own formatting requirement), matching the firm's
    blank chronology template (COMBINED SUMMARY OF MEDICAL RECORDS / TIMELINE — RE/Updated/DOL/
    Facts/DOB/Address/Social-Hx/PMHx/FHx/Surgical-Hx/PCP header block, an Abbreviations/Record
    Sources legend, an Outstanding Records placeholder, then a DATE|PAGE|RECORD|SOURCE|DESCRIPTION
    table) rather than an ad-hoc layout. The browser does the actual review-state bookkeeping
    (sessionStorage — see app.js) and just POSTs the already-assembled, approved-only content here;
    this endpoint is a dumb renderer, not a second source of truth for what's approved. A plain
    single-upload session (no case, no plaintiff name) still gets this exact structure with blank
    placeholders where case-level facts aren't known — same as the firm's own blank template.
    """
    data = request.get_json(force=True, silent=True) or {}
    generated_at = data.get("generated_at") or ""
    total = int(data.get("total") or 0)
    approved = int(data.get("approved") or 0)
    rejected = int(data.get("rejected") or 0)
    pending = total - approved - rejected
    demo = data.get("demographics") or {}

    def field(value, fallback="(not provided)"):
        # _realish() only makes sense for strings ("not stated" etc.) — a numeric value like a
        # page number must pass straight through rather than being treated as a placeholder.
        if value is None or (isinstance(value, str) and not _realish(value)):
            return fallback
        return str(value) or fallback

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # Repeats on every printed/exported page (a real Word header, not just the first line of body
    # text) — confidentiality marking plus live pagination, both standard for a legal document like
    # this. "Page X of Y" uses real field codes (see _add_page_number_field) rather than a computed
    # number, since the actual page count doesn't exist until Word paginates the finished document.
    header_para = doc.sections[0].header.paragraphs[0]
    header_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER — avoids importing the enum for one use
    confidential_run = header_para.add_run(
        f"CONFIDENTIAL - {field(data.get('plaintiff_name'), '(not provided)')}          Page "
    )
    confidential_run.font.name = "Times New Roman"
    confidential_run.font.size = Pt(10)
    confidential_run.bold = True
    _add_page_number_field(header_para, "PAGE")
    of_run = header_para.add_run(" of ")
    of_run.font.name = "Times New Roman"
    of_run.font.size = Pt(10)
    of_run.bold = True
    _add_page_number_field(header_para, "NUMPAGES")

    def heading(text, size=15, bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        return p

    def label_line(label, value):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(field(value))

    heading("COMBINED SUMMARY OF MEDICAL RECORDS / TIMELINE", size=16)
    label_line("RE", data.get("plaintiff_name"))
    label_line("Updated", generated_at)
    label_line("DOL", data.get("dol"))
    label_line("Facts", data.get("facts_summary"))
    if data.get("defendant_names"):
        label_line("Defendant(s)", ", ".join(data["defendant_names"]))
    label_line("DOB", demo.get("dob"))
    label_line("ADDRESS", demo.get("address"))
    label_line("Social Hx", demo.get("social_history"))
    label_line("PMHx", demo.get("past_medical_history"))
    label_line("FHx", demo.get("family_history"))
    label_line("SURGICAL HX", demo.get("surgical_history"))
    label_line("PCP", demo.get("pcp"))

    heading("ABBREVIATIONS AND RECORD SOURCES", size=13)
    record_sources = data.get("record_sources") or []
    if record_sources:
        for entry in record_sources:
            line = _format_record_source_line(
                entry.get("facility"), entry.get("description"), entry.get("date"),
                entry.get("label"), entry.get("page_count"),
            )
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("(not provided)")

    heading("Other record sources/Outstanding records:", size=13)
    doc.add_paragraph(" ")  # blank line for the paralegal to fill in by hand

    heading("Chronological Events:", size=13)
    rows = data.get("rows") or []
    # The column headers are part of the template's fixed skeleton (matching the blank template
    # this is meant to always look like the starting point of) — always shown, even before any
    # chronology entries have been approved, rather than only appearing once there's data.
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    # Default Letter page, default 1" margins each side (never overridden above) => 6.5" usable width.
    # DATE/PAGE/RECORD SOURCE combined ~1/3, DESCRIPTION ~2/3, RECORD SOURCE slightly wider than the
    # other two of that first three — an explicit user requirement, not a stylistic default. Word
    # ignores a plain table.columns[i].width for rows added later via add_row(), so every row's own
    # cells must have their width set individually (see the loop below and inside the row-build loop).
    table.autofit = False
    _COL_WIDTHS = (Inches(0.65), Inches(0.65), Inches(0.85), Inches(4.35))

    def _set_row_widths(cells):
        for cell, width in zip(cells, _COL_WIDTHS):
            cell.width = width

    for col, width in zip(table.columns, _COL_WIDTHS):
        col.width = width
    _set_row_widths(table.rows[0].cells)
    for cell, text in zip(table.rows[0].cells, ("DATE", "PAGE", "RECORD SOURCE", "DESCRIPTION")):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    if rows:
        for r in rows:
            at_issue = bool(r.get("at_issue"))
            cells = table.add_row().cells
            _set_row_widths(cells)

            def set_cell(cell, text, bold=False):
                run = cell.paragraphs[0].add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = bold
                if at_issue:
                    run.font.color.rgb = _AT_ISSUE_RGB

            set_cell(cells[0], field(r.get("date"), ""))
            set_cell(cells[1], field(r.get("page"), ""))
            set_cell(cells[2], field(r.get("source"), ""))
            header_bits = [b for b in (field(r.get("record_type"), ""), field(r.get("author"), "")) if b]
            desc_para = cells[3].paragraphs[0]
            if header_bits:
                bold_run = desc_para.add_run(" -- ".join(header_bits) + "\n")
                bold_run.bold = True
                bold_run.font.name = "Times New Roman"
                bold_run.font.size = Pt(12)
                if at_issue:
                    bold_run.font.color.rgb = _AT_ISSUE_RGB
            body_run = desc_para.add_run(r.get("text") or "")
            body_run.font.name = "Times New Roman"
            body_run.font.size = Pt(12)
            if at_issue:
                body_run.font.color.rgb = _AT_ISSUE_RGB
    else:
        cells = table.add_row().cells
        cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
        cells[0].paragraphs[0].add_run("(no approved chronology entries yet)")

    if data.get("potential_issues"):
        heading("Potential Issues", size=13)
        for item in data["potential_issues"]:
            doc.add_paragraph(item.get("text", ""), style="List Bullet")
            if item.get("citation"):
                p = doc.add_paragraph()
                run = p.add_run(f"Source: {item['citation']}")
                run.italic = True
                run.font.size = Pt(10)

    if data.get("discrepancies"):
        heading("Discrepancies", size=13)
        for item in data["discrepancies"]:
            doc.add_paragraph(item.get("text", ""), style="List Bullet")
            if item.get("citation"):
                p = doc.add_paragraph()
                run = p.add_run(f"Sources: {item['citation']}")
                run.italic = True
                run.font.size = Pt(10)

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Of {total} AI-generated findings: {approved} approved and included above, "
        f"{rejected} rejected and excluded, {pending} not reviewed (also excluded)."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    filename = f"chronology-export-{time.strftime('%Y%m%d-%H%M')}.docx"
    return Response(
        buf.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.route("/session/<session_id>/pdf/<path:filename>")
def serve_pdf(session_id, filename):
    # A session lives in one of two places: the normal single-upload SESSIONS_DIR, or
    # BATCH_SESSIONS_DIR for a group produced by the /case folder feature — case sessions use a
    # "case_<job_id>_<group_key>" id, so they never collide with a normal uuid4().hex session id,
    # but check both roots to serve either kind through the same route (and therefore the same
    # viewer/highlighting code, unmodified).
    for root in (SESSIONS_DIR, BATCH_SESSIONS_DIR):
        candidate = (root / session_id).resolve()
        if root.resolve() in candidate.parents and candidate.is_dir():
            session_dir = candidate
            break
    else:
        return "Session not found or expired", 404

    safe_name = Path(filename).name
    target = session_dir / safe_name
    if not target.is_file():
        fuzzy_match = _find_closest_pdf(session_dir, safe_name)
        if not fuzzy_match:
            return "Not found", 404
        target = fuzzy_match

    return send_from_directory(session_dir, target.name, mimetype="application/pdf")


def _event(payload: dict) -> bytes:
    # direct_passthrough mode (needed for real incremental streaming under Werkzeug's dev
    # server) requires the generator to yield raw bytes rather than str.
    return (json.dumps(payload) + "\n").encode("utf-8")


def _greedy_pack(units: list, budget: int, joiner_len: int) -> list:
    """Shared greedy-packing logic: combines a list of text units into as few groups as possible,
    each group's total length (accounting for the joiner between units) staying within budget.
    Used both to split one oversized document into sub-parts and to pack whole documents into
    chunks — kept as one small function rather than duplicated inline."""
    groups, current, current_len = [], [], 0
    for unit in units:
        unit_len = len(unit) + joiner_len
        if current and current_len + unit_len > budget:
            groups.append(current)
            current, current_len = [], 0
        current.append(unit)
        current_len += unit_len
    if current:
        groups.append(current)
    return groups


def _split_text_to_fit(body: str, budget: int) -> list:
    """Splits one oversized document's body text into sub-parts that each fit within budget.

    Tries progressively finer break points — paragraphs ("\\n\\n"), then lines ("\\n") — and only
    falls back to a raw character-count slice if even single lines don't provide a usable break
    point. Found necessary in practice, not just defensively: this project's own PDF extraction
    joins lines with single newlines throughout (see pdf_to_text/_extract_page_text), so real
    extracted text essentially never contains "\\n\\n" — splitting on paragraphs alone silently
    produced ONE giant unsplit "paragraph" for every real document tested, meaning the very
    oversized-document case this function exists for was never actually being split. Confirmed via
    a direct test against a real 12-page industry-reference sample before this fix existed.
    """
    if len(body) <= budget:
        return [body]

    paragraphs = body.split("\n\n")
    if len(paragraphs) > 1:
        groups = _greedy_pack(paragraphs, budget, joiner_len=2)
        return ["\n\n".join(g) for g in groups]

    lines = body.split("\n")
    if len(lines) > 1:
        groups = _greedy_pack(lines, budget, joiner_len=1)
        return ["\n".join(g) for g in groups]

    # A single line longer than the whole budget (essentially never expected in practice, but
    # must still terminate rather than return one oversized, unsplit part) — raw character slices.
    return [body[i:i + budget] for i in range(0, len(body), budget)]


def _pack_into_chunks(records_parts: list, max_chars: int) -> list:
    """Greedily packs per-document text blocks into chunks that each fit within max_chars,
    keeping whole documents together whenever possible (so the model always sees a complete
    document, never an arbitrary mid-document cut, unless a single document alone exceeds
    max_chars — rare, but real: one of the 7 real industry-reference samples this was tested
    against is 94 pages / ~200K characters on its own, several times over any reasonable context
    budget). In that case the oversized document is itself split at paragraph boundaries into
    sequential sub-parts, each labeled "(part N of M)" in its own header so citations stay
    traceable to which part of the original document a quote came from.

    Returns a list of chunks, each a list of text blocks (each block already includes its own
    "### Document N — filename" header) ready to be joined and sent as one prompt.
    """
    SEPARATOR_LEN = len("\n\n---\n\n")
    expanded = []
    for part in records_parts:
        if len(part) <= max_chars:
            expanded.append(part)
            continue
        # This single document alone exceeds the budget — split it into sub-parts.
        header, _, body = part.partition("\n\n")
        budget = max(max_chars - len(header) - 20, 500)
        sub_parts = _split_text_to_fit(body, budget)
        total = len(sub_parts)
        for i, sub in enumerate(sub_parts, start=1):
            expanded.append(f"{header} (part {i} of {total})\n\n{sub}")

    return _greedy_pack(expanded, max_chars, joiner_len=SEPARATOR_LEN)


_STREAMING_ARRAY_FIELDS = (
    "timeline", "visits", "medications", "procedures", "diagnoses", "labs", "potential_issues", "discrepancies",
)


def _count_ai_findings(findings: dict) -> int:
    """Total AI-generated findings across every streaming category — the "total" half of a group's
    derived review-progress count (see _compute_review_status). Deliberately excludes reviewer-added
    Key Facts; callers add db.load_added_facts()'s count on top, since app.js's own review-progress
    bar (orderedFindingIds) counts both and a server-side total that omitted added facts would
    disagree with what the reviewer actually sees on the review page."""
    return sum(len(findings.get(k) or []) for k in _STREAMING_ARRAY_FIELDS)


def _extract_complete_arrays(partial_json_text: str) -> dict:
    """Best-effort look at an in-progress (not yet valid) JSON object being streamed token-by-token,
    and pulls out whichever top-level array fields are ALREADY fully closed. Deliberately never
    guesses at or force-closes a truncated array/string — the model streams fields in schema order
    (see RESPONSE_SCHEMA), so `"summary"` (last) being unclosed is normal and expected while earlier
    arrays are already complete. A field only comes back here if bracket-matching finds its own `]`
    with a real, independently-parseable JSON array in between — so anything shown mid-stream this
    way is exactly as trustworthy as the same field would be once the whole response finishes.
    """
    results = {}
    for field in _STREAMING_ARRAY_FIELDS:
        m = re.search(rf'"{field}"\s*:\s*\[', partial_json_text)
        if not m:
            continue
        start = m.end() - 1  # index of the opening '['
        depth = 0
        in_string = False
        escape = False
        end = None
        for i in range(start, len(partial_json_text)):
            ch = partial_json_text[i]
            if in_string:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
            elif ch in "[{":
                depth += 1
            elif ch in "]}":
                depth -= 1
                if depth == 0:
                    end = i
                    break
        if end is None:
            continue  # not closed yet — leave it out rather than guess
        try:
            results[field] = json.loads(partial_json_text[start:end + 1])
        except json.JSONDecodeError:
            continue
    return results


def _call_model_for_chunk(records_text: str, detail_level: str, case_context: str = "", on_token_progress=None) -> dict:
    """One call to Ollama for a single (already context-sized) chunk of document text. Raises on a
    request/JSON error — callers decide how to handle a failed chunk.

    Uses Ollama's streaming mode internally (same as the interactive /analyze endpoint) even
    though this itself is a blocking function — not to stream to a browser, but so real, live
    tokens/sec AND whichever findings the model has already fully written can be reported via
    `on_token_progress(tokens_so_far, tokens_per_second, partial_arrays)` while a chunk is still
    generating, rather than only after it completes. A single chunk can take several minutes on
    real documents (measured: ~250-300s is typical) — for a document small enough to fit in one
    chunk, that used to mean a reviewer saw literally nothing until the whole thing finished, even
    though the model streams its JSON fields (timeline, medications, ...) in order as it goes. See
    _extract_complete_arrays for how `partial_arrays` is derived. Throttled to roughly once every
    1.5s so this doesn't hammer the manifest file with a write (or re-scan the response so far) on
    every single token.
    """
    prompt = PROMPT_TEMPLATE.format(
        records=records_text,
        detail_instructions=DETAIL_LEVELS[detail_level],
        case_context=case_context or "No case context was provided for this run.",
    )
    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": True,
        # Restricted to exactly the document ids present in THIS chunk's own text — see
        # _valid_file_ids_in_text/_response_schema_for_ids — so the model cannot reference a
        # document that's in a different chunk or wasn't given to it at all.
        "format": _response_schema_for_ids(_valid_file_ids_in_text(records_text)),
        "options": {"num_ctx": NUM_CTX},
    }

    # Retries a fully degenerate result (every list empty, no summary), OR a request that timed out
    # outright, once before giving up. The degenerate case is verified empirically, not assumed:
    # llama3.1:8b occasionally returns a completely empty-but-schema-valid response for a chunk that
    # plainly has real content to extract — confirmed reproducible by calling the exact same
    # prompt+schema back-to-back and getting a real extraction on one call and nothing on the next.
    # This isn't unique to this app's schema (the pre-existing, much simpler 4-field schema showed
    # the same failure at a lower but nonzero rate too) — it's a real characteristic of this model
    # under grammar-constrained decoding on short documents, and a same-input retry reliably rescues
    # it in practice. See TEST_RESULTS.md. The timeout retry is the same reasoning applied to a
    # different failure mode found in real production use (see TEST_RESULTS.md, 2026-08-03): a slow
    # machine hitting the per-call timeout isn't necessarily stuck, just having a slow moment
    # (thermal throttling, a background process competing for CPU) — one retry gives it a genuine
    # second chance before the whole patient group gets marked failed.
    max_attempts = 2
    for attempt in range(1, max_attempts + 1):
        start = time.time()
        full_response = ""
        tokens_seen = 0
        eval_count = 0
        eval_duration_s = 0
        last_report_time = start

        try:
            # Long timeout — this runs from an unattended background thread overnight, not a live
            # browser request, so there's no UX reason to time out quickly the way an interactive
            # request might. Was 1800s (30 min); raised after a real production timeout on a slower
            # machine — see the retry comment above and NUM_CTX's own comment, both from the same
            # incident. Chunk size was also lowered at the same time (via NUM_CTX), so this is a
            # backstop for occasional slow calls, not the primary mitigation.
            with requests.post(OLLAMA_URL, json=payload, stream=True, timeout=3600) as resp:
                resp.raise_for_status()
                for line in resp.iter_lines():
                    if not line:
                        continue
                    chunk = json.loads(line)
                    full_response += chunk.get("response", "")

                    if chunk.get("done"):
                        eval_count = chunk.get("eval_count", tokens_seen)
                        eval_duration_s = chunk.get("eval_duration", 0) / 1e9
                        break

                    tokens_seen += 1
                    now = time.time()
                    if on_token_progress and (now - last_report_time) >= 1.5:
                        elapsed = now - start
                        partial_arrays = _extract_complete_arrays(full_response)
                        on_token_progress(tokens_seen, tokens_seen / elapsed if elapsed > 0 else 0, partial_arrays)
                        last_report_time = now
        except requests.exceptions.Timeout:
            if attempt < max_attempts:
                applog.log_event(
                    "chunk_timeout_retry",
                    f"Chunk timed out on attempt {attempt} of {max_attempts}, retrying once",
                    level="WARNING",
                )
                continue
            raise

        wall_time = time.time() - start
        findings = json.loads(full_response)  # raises JSONDecodeError -> caller marks failed
        if attempt < max_attempts and _looks_degenerate(findings):
            continue
        break

    return {"findings": findings, "wall_time": wall_time, "eval_count": eval_count, "eval_duration_s": eval_duration_s}


_DATED_LIST_FIELDS = ("timeline", "visits", "medications", "procedures", "diagnoses", "labs")

# Every format actually seen coming out of real medical records/EHR exports so far — tried in
# order, first match wins. Deliberately a fixed list rather than a lenient general-purpose parser
# (python-dateutil-style guessing): an unrecognized format is left completely alone rather than
# risk silently misreading it, same "omit rather than guess" principle the prompts already use for
# the model's own extraction.
_DATE_INPUT_FORMATS = (
    "%m/%d/%Y", "%m/%d/%y",                              # already slash-separated
    "%m-%d-%Y", "%m-%d-%y",                              # dash-separated, US order
    "%Y-%m-%d", "%y-%m-%d",                              # ISO, sometimes seen in EHR exports
    "%B %d, %Y", "%B %d %Y", "%B %d, %y", "%B %d %y",    # "March 4, 2026"
    "%b %d, %Y", "%b %d %Y", "%b %d, %y", "%b %d %y",    # "Mar 4, 2026"
    "%d %B %Y", "%d %b %Y",                              # "4 March 2026"
    "%d-%b-%Y", "%d-%B-%Y",                              # "04-Mar-2026"
)


def _normalize_date_str(raw: str) -> str:
    """Best-effort reformat to this app's standard chronology date format, MM/DD/YY (e.g.
    "03/04/26") — always slash-separated, never dashes. Deterministic and code-side, not a prompt
    instruction: the model is asked to extract each date verbatim from its source (needed for exact
    citation matching), and reformatting is exactly the kind of transformation step that's invited
    hallucination before (see the prompts' own "don't invent a clock time" note) — so this is a
    fixed, auditable parse-and-reformat pass applied afterward, not something asked of the model.
    Returns `raw` unchanged if it doesn't cleanly match one of _DATE_INPUT_FORMATS.
    """
    raw = (raw or "").strip()
    if not raw:
        return raw
    for fmt in _DATE_INPUT_FORMATS:
        try:
            parsed = datetime.strptime(raw, fmt)
        except ValueError:
            continue
        return parsed.strftime("%m/%d/%y")
    return raw


def _normalize_dates(findings: dict) -> None:
    """Mutates `findings` in place — every dated-entry list's "date" field reformatted via
    _normalize_date_str. Called once per pipeline right alongside the other post-processing passes
    (_resolve_bates/_dedupe_all_with_multi_source/_assign_finding_ids) — safe to run at any point
    relative to those, since none of them key off of "date" (Bates resolution and dedup both key off
    "quote"/"text", never "date")."""
    for key in _DATED_LIST_FIELDS:
        for item in findings.get(key, []) or []:
            item["date"] = _normalize_date_str(item.get("date", ""))


def _looks_degenerate(findings: dict) -> bool:
    """True if every dated-entry list AND the summary came back empty — a strong signal of the
    schema-constrained-generation failure mode described above, not a legitimately empty document
    (a real document with nothing extractable would still normally get at least a one-line
    summary noting that)."""
    return not any(findings.get(key) for key in _DATED_LIST_FIELDS) and not (findings.get("summary") or "").strip()


# Matches the per-page marker pdf_to_text always embeds — group 1 is the page number (always
# present), group 2 is the Bates number (only present/captured when that page is actually stamped).
_PAGE_MARKER_PATTERN = re.compile(r"\[PAGE (\d+)\]\n(?:\[BATES: ([^\]]+)\]\n)?")


def _normalize_quote_text(s: str) -> str:
    """Shared by _find_citation_for_quote and _reconcile_source_files — normalizes the same two
    real, observed variant classes (smart quotes/dashes, collapsed whitespace) so a quote match
    isn't defeated by a cosmetic difference between the model's reproduction and the source text."""
    for orig, repl in (("‘", "'"), ("’", "'"), ("“", '"'), ("”", '"'),
                       ("–", "-"), ("—", "-")):
        s = s.replace(orig, repl)
    return re.sub(r"\s+", " ", s.strip())


def _resolve_file_ids(findings: dict, file_id_to_filename: dict) -> None:
    """Converts every finding's model-reported `file_id` (a small integer — see _dated_entry_schema)
    back into a real `source_file` filename string, using the exact same id->filename map the
    document headers were built from (see _generate_chronology/the /analyze route). Must run FIRST,
    before _reconcile_source_files/_resolve_bates/anything else — once this runs, every downstream
    consumer of `source_file`/`source_files` is completely unchanged from before file_id existed.

    An id the map doesn't recognize (not yet possible to fully prevent at this stage — the enum
    constraint that makes it structurally impossible is a later addition) resolves to None, the same
    "honestly unresolvable" value every downstream consumer (_resolve_bates, PDF serving, the Record
    Sources legend) already knows how to handle gracefully — never guessed/invented.
    """
    def to_filename(file_id) -> str | None:
        return file_id_to_filename.get(file_id) if file_id is not None else None

    for key in _DATED_LIST_FIELDS + ("potential_issues",):
        for item in findings.get(key, []) or []:
            item["source_file"] = to_filename(item.pop("file_id", None))
    for item in findings.get("discrepancies", []) or []:
        item["source_files"] = [to_filename(fid) for fid in (item.pop("file_ids", None) or [])]
    for item in findings.get("record_sources", []) or []:
        item["source_file"] = to_filename(item.pop("file_id", None))


def _reconcile_source_files(findings: dict, filename_to_text: dict) -> None:
    """Corrects a finding's source_file when it doesn't match any real ingested document — a real,
    observed failure mode on larger real-world cases (confirmed via user report: a click sometimes
    rendered a real but WRONG PDF, or 404'd entirely): the model occasionally cites the wrong
    filename, a genuinely different real file rather than just a punctuation/case variant
    _find_closest_pdf already tolerates client-side for PDF serving. Recovery, in this order, never
    guessing if neither works:
    1. The finding's own quote is required to be verbatim (enforced elsewhere in the prompt) — if
       that exact quote is found in exactly one OTHER real document's text, that document is
       unambiguously the true source regardless of what filename the model wrote down. Ambiguous
       (found in 0 or 2+ documents) is left alone rather than guessed.
    2. Otherwise, fall back to the same fuzzy filename normalization _find_closest_pdf already uses
       (case/punctuation/whitespace-insensitive) — catches a near-miss reproduction of a real name.
    If neither recovers a real match, source_file is left exactly as the model wrote it — downstream
    (bates resolution, PDF serving, the Record Sources legend) already treat an unresolvable
    source_file as an honest "can't verify this citation" rather than silently guessing, and this
    function must not undermine that by forcing a match where none is truly warranted. Must run
    BEFORE _resolve_bates/_resolve_record_sources so every later consumer sees the corrected name.

    IMPORTANT: the quote check runs UNCONDITIONALLY, even when `src` already looks like a real,
    recognized filename — do not reintroduce a `src in filename_to_text: return src` short-circuit
    ahead of it. Once file_id->filename conversion (_resolve_file_ids) runs first, `source_file` is
    ALWAYS a real filename by construction (the model can no longer hallucinate an arbitrary
    string) — a short-circuit on "already a real filename" would then return true for every single
    finding and silently disable this function's actual remaining job: catching a syntactically
    valid file_id that the model nonetheless attributed to the WRONG real document. The quote is the
    only signal that can still catch that case, so it must always get a chance to override `src`.
    """
    def reconcile_one(src: str | None, quote: str | None) -> str | None:
        if not src:
            return src
        if quote:
            norm_quote = _normalize_quote_text(quote)
            if norm_quote:
                matches = [
                    fname for fname, text in filename_to_text.items()
                    if norm_quote in _normalize_quote_text(text)
                ]
                if len(matches) == 1:
                    return matches[0]
        if src in filename_to_text:
            return src
        normalized_src = _normalize_filename(src)
        if normalized_src:
            for fname in filename_to_text:
                if _normalize_filename(fname) == normalized_src:
                    return fname
        return src

    for key in _DATED_LIST_FIELDS + ("potential_issues",):
        for item in findings.get(key, []) or []:
            item["source_file"] = reconcile_one(item.get("source_file"), item.get("quote"))
    for item in findings.get("discrepancies", []) or []:
        source_files = item.get("source_files") or []
        quotes = item.get("quotes") or []
        item["source_files"] = [
            reconcile_one(src, quotes[i] if i < len(quotes) else None)
            for i, src in enumerate(source_files)
        ]


def _find_citation_for_quote(doc_text: str | None, quote: str | None) -> tuple[str | None, int | None]:
    """Given a document's full extracted text (with a `[PAGE N]` marker — and, when the page is
    actually Bates-stamped, a following `[BATES: X]` marker — embedded once per page, see
    pdf_to_text) and a quote the model claims came from this document, finds which page actually
    contains that quote and returns `(bates, page_number)`. `bates` is None whenever that specific
    page has no genuine Bates stamp (never a fabricated placeholder — a real, previously-existing
    bug: pages with no stamp used to get a literal "UNKNOWN" string that flowed through as if it
    were real); `page_number` is the 1-indexed page within THIS document, always present when the
    quote is found at all, and exists specifically so the caller has a real fallback citation when
    a document simply isn't Bates-stamped, rather than nothing at all.

    Deliberately NOT a schema field the model reports itself — both values are derived server-side
    from the quote the model already has to provide (verbatim, strictly grounded), the same way
    highlight positions are independently verified rather than trusted to the model's own claim
    elsewhere in this app. This also sidesteps a real, measured risk: adding another required
    per-entry schema field was what caused the empty-output reliability regression documented above
    — deriving this instead of asking the model to report it adds the capability with zero added
    schema complexity.

    Whitespace- and punctuation-normalized substring match, not a raw exact one — confirmed
    necessary via direct testing against a real multi-document case: (1) a quote can have a newline
    or collapsed run of spaces where the source text has something slightly different, and (2) the
    model sometimes reproduces a quote's punctuation marks in a different but equivalent style (a
    curly/smart quote where the source has a straight one, or vice versa, similarly for dashes) —
    the exact same two failure classes the client-side highlight search already has to tolerate
    (see findQuoteInDoc/normalize in app.js). A naive exact-substring check silently failed to
    resolve a Bates number for several genuinely-correct quotes because of this. Case-sensitive —
    only whitespace and this specific set of punctuation-style variants are normalized, to avoid
    over-matching against unrelated text.
    """
    if not doc_text or not quote:
        return None, None
    markers = list(_PAGE_MARKER_PATTERN.finditer(doc_text))
    if not markers:
        return None, None

    norm_quote = _normalize_quote_text(quote)
    if not norm_quote:
        return None, None
    for i, m in enumerate(markers):
        seg_start = m.end()
        seg_end = markers[i + 1].start() if i + 1 < len(markers) else len(doc_text)
        norm_segment = _normalize_quote_text(doc_text[seg_start:seg_end])
        if norm_quote in norm_segment:
            return m.group(2), int(m.group(1))
    return None, None


def _resolve_bates(findings: dict, filename_to_text: dict) -> None:
    """Mutates `findings` in place, attaching a derived `bates`/`page_number` (or `bates_list`/
    `page_number_list` for discrepancies, which cite multiple source files at once) to every citable
    item — see _find_citation_for_quote. A quote that can't be located in its claimed source (e.g. a
    near-miss paraphrase, or a source file whose text couldn't be read) gets both as None rather
    than a guessed/fabricated value — the frontend treats that as any other unresolvable citation,
    never inventing a page reference. `page_number` exists specifically as a real fallback for the
    PAGE column when a document simply isn't Bates-stamped (bates is None but the quote genuinely
    was located, so its page IS known). Applied to every entry type this app produces a citation
    for: the five dated-entry lists, potential_issues, and discrepancies.
    """
    for key in _DATED_LIST_FIELDS + ("potential_issues",):
        for item in findings.get(key, []) or []:
            doc_text = filename_to_text.get(item.get("source_file"))
            item["bates"], item["page_number"] = _find_citation_for_quote(doc_text, item.get("quote"))
    for item in findings.get("discrepancies", []) or []:
        source_files = item.get("source_files") or []
        quotes = item.get("quotes") or []
        resolved = [
            _find_citation_for_quote(filename_to_text.get(src), quote)
            for src, quote in zip(source_files, quotes)
        ]
        item["bates_list"] = [b for b, _ in resolved]
        item["page_number_list"] = [p for _, p in resolved]


def _dedupe_with_multi_source(items: list) -> list:
    """Merges items whose "text" is an exact-normalized match into ONE entry with every distinct
    source it came from, rather than either showing the same event/fact once per duplicate source
    (real case files commonly contain literal 1:1 duplicate records — the same visit obtained both
    directly from a provider and as a "PL" copy, or simply the same document uploaded twice) or
    silently dropping the duplicates (losing the fact that multiple records independently
    corroborate it, which is itself useful information — e.g. a "PL" copy and a certified vendor
    copy both stating the same thing). Order-preserving: the merged entry sits at the position of
    its first occurrence.

    Must run AFTER _resolve_bates — relies on each raw item's own already-resolved "bates" value to
    build the merged item's "bates_list", so this can't happen any earlier in the pipeline.
    """
    groups: dict = {}
    blank_counter = 0
    for item in items:
        key = re.sub(r"\s+", " ", (item.get("text") or "").strip().lower())
        if not key:
            blank_counter += 1
            key = f"__blank_{blank_counter}__"  # never merge with anything else
        groups.setdefault(key, []).append(item)

    merged = []
    for group in groups.values():
        first = dict(group[0])
        if len(group) > 1:
            seen_files = []
            bates_list = []
            for g in group:
                sf = g.get("source_file")
                if sf and sf not in seen_files:
                    seen_files.append(sf)
                    bates_list.append(g.get("bates"))
            if len(seen_files) > 1:
                first["source_files"] = seen_files
                first["bates_list"] = bates_list
        merged.append(first)
    return merged


def _dedupe_all_with_multi_source(findings: dict) -> None:
    """Applies _dedupe_with_multi_source to every entry-type list this app produces a citation
    for — the five dated-entry lists plus potential_issues (discrepancies are deliberately excluded:
    they already have their own genuine multi-source shape from the model directly, describing a
    conflict BETWEEN sources, which is a different thing from the same fact being corroborated by
    duplicates). Mutates findings in place. Must run after _resolve_bates."""
    for key in _DATED_LIST_FIELDS + ("potential_issues",):
        findings[key] = _dedupe_with_multi_source(findings.get(key) or [])


_FINDING_ID_SECTIONS = ("timeline", "potential_issues", "discrepancies")


def _assign_finding_ids(findings: dict) -> None:
    """Stamps a stable "finding_id" onto every item in timeline/potential_issues/discrepancies —
    the ONLY sections the frontend renders as individually reviewable/approvable rows from
    server-provided data (medications/procedures/diagnoses/labs are Key-Facts-tab source data only;
    key_facts items are reviewer-added, already keyed by their own sequential id). Computed from
    CONTENT (section + source file(s) + quote/text), not array position, so the same item gets the
    same id whether seen in a live partial-generation snapshot or the final post-dedup result —
    critical for durable review state (see db.py's review_actions table) to survive dedup safely,
    and for a previously-approved item's status to still apply after incremental reprocessing
    appends new items and re-merges these same lists (see _merge_incremental_group_result).

    Previously, the frontend keyed potential_issues/discrepancies by raw array position
    (`${sectionKey}-${index}` in app.js's idFor) — the same reassign-on-reorder bug already fixed
    for timeline specifically (see "Chronological ordering..." in TEST_RESULTS.md), just not yet
    closed for these other two sections. This closes it for all three, server-side, once.

    Mutates `findings` in place. Idempotent and safe to call repeatedly on growing/re-merged data —
    re-running it on an already-stamped list recomputes the same ids for unchanged items.
    """
    for section in _FINDING_ID_SECTIONS:
        items = findings.get(section) or []
        seen: dict[str, int] = {}
        for item in items:
            source_key = item.get("source_file") or "|".join(item.get("source_files") or [])
            quote_key = item.get("quote") or "|".join(item.get("quotes") or [])
            raw = f"{section}:{source_key}:{quote_key}:{item.get('text', '')}"
            base_id = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]
            n = seen.get(base_id, 0)
            item["finding_id"] = base_id if n == 0 else f"{base_id}-{n + 1}"
            seen[base_id] = n + 1


def _cap_label_words(label: str, max_words: int = 2) -> str:
    """Enforces the "2 words max" rule server-side rather than trusting the model's own compliance
    (same "don't trust the model's own claim, enforce deterministically" reasoning already used for
    uniqueness below) — collapses repeated whitespace to single spaces and truncates to the first
    `max_words` words."""
    words = re.sub(r"\s+", " ", label.strip()).split(" ")
    return " ".join(w for w in words[:max_words] if w)


def _slugify_label(filename: str) -> str:
    """Fallback label derived purely from a filename — used when the model didn't cover a source
    file at all (or hallucinated one that doesn't match anything real). Title-cased, non-alphanumeric
    characters treated as word boundaries, capped at 2 words to match the style asked of the model
    (e.g. "op_report_v2.pdf" -> "Op Report")."""
    stem = Path(filename).stem
    words = [w.capitalize() for w in re.split(r"[^A-Za-z0-9]+", stem) if w]
    return " ".join(words[:2]) or "Record"


def _resolve_record_sources(findings: dict, source_filenames: list) -> dict:
    """Returns {filename: short_label} covering every actually-ingested source file (falls back to
    a filename-derived label for any the model skipped or hallucinated a mismatched filename for)
    and DETERMINISTICALLY enforces uniqueness — appending "2", "3", ... on a collision — rather
    than trusting the model's own claim that its labels are already unique, since correctness of
    uniqueness here is a hard requirement (two different source files must never share a label),
    not a best-effort nicety.
    """
    model_labels = {}
    for item in findings.get("record_sources", []) or []:
        src = item.get("source_file")
        label = _cap_label_words(item.get("label") or "")
        if src in source_filenames and label:
            model_labels[src] = label

    labels = {}
    seen_lower = set()
    for src in source_filenames:
        base = model_labels.get(src) or _slugify_label(src)
        candidate = base
        n = 2
        while candidate.lower() in seen_lower:
            candidate = f"{base} {n}"
            n += 1
        seen_lower.add(candidate.lower())
        labels[src] = candidate
    return labels


def _format_record_source_line(facility, description, date, label, page_count) -> str:
    """One line for the Abbreviations/Record Sources legend: "{Facility} – {Description} {date} –
    ({Abbreviation}) – {N page(s)}" — e.g. "Beaumont – Canton CT Temporal Report 7.27.21 – (CT Temp)
    – 2 pages". Unlike the fixed-slot chronology table, a legend line has no fixed shape to
    preserve — any missing optional piece (facility/description/date, all best-effort since not every
    document clearly states them) is skipped entirely rather than shown as a "(not provided)"
    filler, so a sparse document still reads as a normal, complete-looking line rather than a
    template with visible gaps."""
    lead_bits = [b for b in (facility, description, date) if b and _realish(b)]
    parts = []
    if lead_bits:
        parts.append(" ".join(lead_bits))
    if label and _realish(label):
        parts.append(f"({label})")
    if page_count:
        parts.append(f"{page_count} page{'s' if page_count != 1 else ''}")
    return " – ".join(parts) if parts else "(not provided)"


# Label variants to grep for per patient_demographics subfield — see _verify_demographics_with_grep.
_DEMOGRAPHIC_LABEL_PATTERNS = {
    "dob": ["Date of Birth", "DOB"],
    "address": ["Address"],
    "social_history": ["Social History", "Social Hx", "Soc Hx"],
    "past_medical_history": ["Past Medical History", "PMHx", "PMH"],
    "family_history": ["Family History", "FHx", "FH"],
    "surgical_history": ["Past Surgical History", "Surgical History", "PSH", "Surgical Hx"],
    "pcp": ["Primary Care Physician", "Primary Care Provider", "PCP"],
}


def _grep_after_label(text: str, labels: list) -> str | None:
    """Looks for 'Label: value' or 'Label - value' and returns the text immediately after the
    label, stopping at the first line break or blank-line boundary — a plain text search, not
    another model call. Best-effort: a safety net for the case where the model reported "not
    stated" but the information was directly there in a labeled field the model's own read simply
    missed, not a substitute for actually reading the document."""
    for label in labels:
        m = re.search(rf"{re.escape(label)}\s*[:\-]\s*(.+)", text, re.IGNORECASE)
        if m:
            value = m.group(1).strip().splitlines()[0].strip()
            if value and _realish(value):
                return value
    return None


def _verify_demographics_with_grep(findings: dict, combined_text: str) -> None:
    """Safety net specifically for patient_demographics: before the export ever declares a field
    "not provided," double-check it against a direct text search of everything actually extracted
    from the case's documents — catches the model reporting "not stated" for something that IS
    literally present in a labeled field it simply missed while generating a long, complex
    response. Never overrides a value the model DID find (that's trusted as-is, same as
    everywhere else in this app); only fills in fields still left as a placeholder. Anything found
    this way is marked so a reviewer knows it came from a plain text search, not the model's own
    read, and should still be double-checked against the source.
    """
    demo = findings.setdefault("patient_demographics", {})
    for field, labels in _DEMOGRAPHIC_LABEL_PATTERNS.items():
        if _realish(demo.get(field)):
            continue  # model already found something real — trust it, don't second-guess it
        found = _grep_after_label(combined_text, labels)
        if found:
            demo[field] = f"{found} (found via text search — verify against source)"


# Same placeholder-value convention as the frontend's PLACEHOLDER_VALUES/realValue() in app.js —
# the prompt tells the model to write "not stated" rather than guess, so this filters those
# placeholders out server-side wherever a "did we actually learn something" check matters (e.g.
# picking which chunk's patient_demographics value to keep — see _merge_chunk_findings).
_PLACEHOLDER_VALUES = {"", "not stated", "not applicable", "n/a", "na", "none", "unknown", "not specified", "not given"}


def _realish(value) -> bool:
    return isinstance(value, str) and value.strip().lower() not in _PLACEHOLDER_VALUES


_ATOMIC_LIST_FIELDS = ("medications", "procedures", "diagnoses", "labs", "potential_issues")


def _merge_chunk_findings(chunk_findings: list) -> dict:
    """Combines N chunks' worth of independently-generated findings into one result. Each chunk
    was analyzed with no visibility into the others, so this is a straightforward concatenation for
    every field (order preserved — chunk order follows document order, same as the model's own
    within-chunk ordering already relied on). Deduplicating exact-repeat facts across chunks (e.g.
    the same medication mentioned in two different chunks, or a literal duplicate source document
    landing in a different chunk than its twin) happens SEPARATELY, later, in
    _dedupe_all_with_multi_source — that step needs each item's own resolved Bates number first
    (see _resolve_bates), which doesn't exist yet at merge time, so it can't happen here.

    Explicitly NOT attempted: cross-chunk discrepancy detection, or noticing a trend that only
    becomes visible spanning multiple chunks (e.g. a gradual lab-value drift split across chunk
    boundaries) — a real, known limitation of chunking, not silently pretended away. See the
    warning this triggers in _generate_chronology.
    """
    merged = {"timeline": [], "visits": [], "medications": [], "procedures": [], "diagnoses": [], "labs": [],
              "potential_issues": [], "discrepancies": [], "summary": [], "patient_demographics": {},
              "record_sources": []}
    for cf in chunk_findings:
        # record_sources: each chunk only ever sees a disjoint subset of the case's documents (a
        # document is never split across chunks unless individually oversized — see
        # _pack_into_chunks), so plain concatenation is correct here, same as timeline/visits/discrepancies.
        for key in ("timeline", "visits", "discrepancies", "record_sources"):
            merged[key].extend(cf.get(key, []))
        for key in _ATOMIC_LIST_FIELDS:
            merged[key].extend(cf.get(key, []))
        if cf.get("summary"):
            merged["summary"].append(cf["summary"])
        # Case-wide, not per-chunk — first chunk to actually state a subfield wins. A documented
        # simplification (not an adjudication between conflicting chunks), same spirit as the rest
        # of this merge: good enough for a quick-reference header, not a substitute for reviewing
        # the source documents directly.
        for subfield, value in (cf.get("patient_demographics") or {}).items():
            if _realish(value) and not merged["patient_demographics"].get(subfield):
                merged["patient_demographics"][subfield] = value

    merged["summary"] = "\n\n".join(
        f"[Part {i + 1}] {s}" for i, s in enumerate(merged["summary"])
    ) if len(merged["summary"]) > 1 else (merged["summary"][0] if merged["summary"] else "")

    return merged


def _generate_chronology(pdf_paths: list, detail_level: str, case_context: str = "", on_progress=None,
                          file_ids: dict[str, int] | None = None) -> dict:
    """Blocking (non-streaming) version of the /analyze pipeline's extraction + generation logic —
    used by the batch-folder feature, where a background thread processes many patient groups
    unattended and there's no live browser tab to stream progress events to. Raises ValueError if
    nothing extractable was found, or lets a requests/JSON error propagate — the caller (the batch
    job loop) is responsible for catching either and marking that specific patient group failed
    without aborting the rest of the overnight run.

    Handles a patient group whose combined document text exceeds the model's context window by
    splitting it into multiple sequential chunks (see _pack_into_chunks) and merging each chunk's
    independently-generated findings (see _merge_chunk_findings) — found necessary in practice
    (not a hypothetical): a single 12-page real-world document already exceeded the default
    context budget once the prompt's own instructional overhead is accounted for, and this
    feature's whole purpose is to reliably process cases far larger than that. This replaces the
    previous behavior of silently truncating anything past the context limit, which discarded
    content with no way to recover it.

    `on_progress`, if given, is called as `on_progress(status_text, partial_findings)` at
    meaningful points — after each document is read, and after each chunk finishes generating
    (with the merge-so-far of every chunk completed up to that point, so the caller can show a
    real, growing chronology while later chunks are still running rather than nothing at all until
    the entire — potentially very long, multi-chunk — analysis finishes).

    `file_ids`, if given (Case Mode only — see db.get_or_assign_file_ids), maps each pdf_path's own
    filename to a DURABLE per-job integer id that survives across this case's initial run and every
    later rescan, used as the "### Document N — filename" header's number instead of this call's own
    positional index. Single-upload mode (never rescanned, one-shot) leaves this None and gets
    today's simple positional numbering — no cross-call stability needed there.
    """
    records_parts = []
    warnings = []
    ocr_files = []
    source_filenames = []
    filename_to_text = {}  # for _resolve_bates below — the exact per-file text the model actually saw
    file_id_to_filename = {}  # for _resolve_file_ids — converts the model's file_id back to a real filename
    page_counts = {}  # for the Record Sources legend's "N page(s)" — a hard fact, never model-derived
    total_pages = 0

    for i, pdf_path in enumerate(pdf_paths, start=1):
        doc_id = (file_ids or {}).get(pdf_path.name, i)
        if on_progress:
            on_progress(f"Reading document {i} of {len(pdf_paths)} ({pdf_path.name})…", None)

        try:
            text, num_pages, num_ocr_pages = pdf_to_text(pdf_path)
        except Exception as e:
            warnings.append(f"Could not read {pdf_path.name}: {e}")  # shown to the reviewer who
            # uploaded this exact file — their own filename back to them isn't a PII leak, unlike
            # persisting it to the log file below, which deliberately omits it.
            applog.log_event(
                "document_extraction_failed", f"Failed to extract text from a document: {type(e).__name__}",
                level="ERROR", context={"document_index": i, "document_count": len(pdf_paths)},
                exc_info=sys.exc_info(),
            )
            continue

        total_pages += num_pages
        clean_text = text.strip()

        if num_ocr_pages:
            ocr_files.append(pdf_path.name)
            warnings.append(
                f"'{pdf_path.name}': {num_ocr_pages} of {num_pages} page(s) had no usable text "
                "layer (likely scanned/faxed) and were read via OCR instead — double-check "
                "those findings against the source image, since OCR is less reliable than a "
                "native text layer."
            )

        if not clean_text:
            warnings.append(
                f"'{pdf_path.name}': no extractable text found even after OCR "
                f"({num_pages} page(s)) — this file was excluded from the analysis."
            )
            continue

        records_parts.append(f"### Document {doc_id} — {pdf_path.name}\n\n{clean_text}")
        source_filenames.append(pdf_path.name)
        filename_to_text[pdf_path.name] = clean_text
        file_id_to_filename[doc_id] = pdf_path.name
        page_counts[pdf_path.name] = num_pages

    if not records_parts:
        raise ValueError("No extractable text found in any document for this patient group.")

    chunks = _pack_into_chunks(records_parts, MAX_INPUT_CHARS)

    if len(chunks) > 1:
        warnings.append(
            f"This case's documents totaled more than the model can analyze in a single pass and "
            f"were processed in {len(chunks)} sequential chunks instead of being truncated — every "
            "document was fully analyzed, but pattern/discrepancy detection that would require "
            "seeing two chunks' content at the same time (e.g. a gradual trend split across chunk "
            "boundaries) may be less reliable than a single-pass analysis. Review carefully."
        )

    chunk_results = []
    total_eval_count = 0
    total_eval_duration_s = 0.0
    total_wall_time = 0.0
    for chunk_num, chunk_parts in enumerate(chunks, start=1):
        pass_label = f"pass {chunk_num} of {len(chunks)}" if len(chunks) > 1 else "generating"
        if on_progress:
            on_progress(f"Generating chronology ({pass_label})…", None)

        def on_token_progress(tokens_so_far, tokens_per_second, partial_arrays, pass_label=pass_label):
            if not on_progress:
                return
            status_text = (
                f"Generating chronology ({pass_label}) — {tokens_so_far} tokens so far, "
                f"{tokens_per_second:.1f} tok/s"
            )
            # Convert file_id -> source_file on this still-streaming partial parse too, same as the
            # completed-chunk result below — otherwise a live mid-generation preview would briefly
            # show raw integer file_ids instead of filenames before the chunk actually finishes.
            _resolve_file_ids(partial_arrays, file_id_to_filename)
            # Merge whatever earlier chunks already fully finished with whatever fields THIS
            # chunk has already fully written, so a single-chunk case (the common one — a lone
            # small-to-medium patient file) still gets a growing chronology mid-generation instead
            # of showing nothing at all until the entire chunk's JSON is done. Real findings only —
            # see _extract_complete_arrays, which never guesses at a still-truncated array.
            if any(partial_arrays.values()):
                partial = _merge_chunk_findings(chunk_results + [partial_arrays])
                _assign_finding_ids(partial)
                on_progress(status_text, partial)
            else:
                on_progress(status_text, None)

        chunk_text = "\n\n---\n\n".join(chunk_parts)
        result = _call_model_for_chunk(
            chunk_text, detail_level, case_context=case_context, on_token_progress=on_token_progress
        )
        # Converted here, right as each chunk's OWN result comes back — every later step (the live
        # progress merge above, the final cross-chunk merge below, _reconcile_source_files/
        # _resolve_bates) works with real source_file strings only and needs no file_id awareness.
        _resolve_file_ids(result["findings"], file_id_to_filename)
        chunk_results.append(result["findings"])
        total_eval_count += result["eval_count"]
        total_eval_duration_s += result["eval_duration_s"]
        total_wall_time += result["wall_time"]

        if on_progress:
            partial = _merge_chunk_findings(chunk_results) if len(chunk_results) > 1 else chunk_results[0]
            _assign_finding_ids(partial)
            chunk_tok_s = round(result["eval_count"] / result["eval_duration_s"], 1) if result["eval_duration_s"] else None
            speed_note = f" ({chunk_tok_s} tok/s)" if chunk_tok_s else ""
            done_label = f"{chunk_num} of {len(chunks)} passes done{speed_note} — chronology below updates as more complete" \
                if len(chunks) > 1 else f"Finalizing…{speed_note}"
            on_progress(done_label, partial)

    findings = _merge_chunk_findings(chunk_results) if len(chunk_results) > 1 else chunk_results[0]
    _normalize_dates(findings)
    _reconcile_source_files(findings, filename_to_text)
    _resolve_bates(findings, filename_to_text)
    _dedupe_all_with_multi_source(findings)
    _assign_finding_ids(findings)  # after dedup, so a merged item's id is computed on its final
                                    # content, not a stale pre-merge id — see _assign_finding_ids
    record_source_labels = _resolve_record_sources(findings, source_filenames)
    _verify_demographics_with_grep(findings, "\n\n".join(filename_to_text.values()))

    return {
        "findings": findings,
        "stats": {
            "documents_processed": len(records_parts),
            "total_pages": total_pages,
            "wall_time_seconds": round(total_wall_time, 1),
            "output_tokens": total_eval_count,
            "tokens_per_second": round(total_eval_count / total_eval_duration_s, 1) if total_eval_duration_s else None,
            "input_truncated": False,  # chunking replaces truncation — nothing is silently discarded
            "chunks": len(chunks),
            "model": MODEL,
            "detail_level": detail_level,
        },
        "warnings": warnings,
        "ocr_files": ocr_files,
        "source_filenames": source_filenames,
        "record_source_labels": record_source_labels,
        "page_counts": page_counts,
    }


def _copy_files_into_session(session_dir: Path, files: list) -> list:
    """Copies `files` into `session_dir` (creating it if needed) and returns the copied paths —
    shared by both initial case processing (Phase 4 below) and incremental rescans
    (_run_case_rescan), since a rescan adds files to an EXISTING session dir without disturbing
    files already copied there for earlier findings' highlighting to keep working."""
    session_dir.mkdir(parents=True, exist_ok=True)
    copied_paths = []
    for src in files:
        dest = session_dir / src.name
        shutil.copy2(src, dest)
        copied_paths.append(dest)
    return copied_paths


_UNSAFE_FILENAME_CHARS = re.compile(r'[\x00-\x1f\\/:*?"<>|]')


def _sanitize_filename_segment(name: str) -> str:
    """Strips only characters that are genuinely unsafe in a filesystem path segment (control
    characters, path separators, and the handful of characters Windows disallows in a filename) —
    deliberately NOT Werkzeug's secure_filename(), which is far more aggressive than path-traversal
    safety actually requires and replaces spaces (along with most punctuation) with underscores.
    Real medical record filenames commonly contain spaces and carry real information in them (e.g.
    "Beaumont - Canton CT Temporal Report 7.27.21.pdf") that a mangled underscore-name loses for
    display purposes (record source labels, the abbreviations legend) — preserving them here doesn't
    weaken path-traversal safety, which _safe_relative_path already enforces separately by stripping
    ".."/"." segments before this ever runs."""
    cleaned = _UNSAFE_FILENAME_CHARS.sub("", name).strip()
    return cleaned or "_"


def _safe_relative_path(raw_path: str) -> Path | None:
    """Sanitizes a browser-supplied relative path (a webkitRelativePath, e.g.
    "case_smith/pleadings/complaint.pdf", arriving as an uploaded file's filename) into something
    safe to join under a server-controlled directory. Rejects/strips ".."/"." segments and absolute
    roots so a malicious filename can never escape CASE_SOURCES_DIR, but otherwise preserves the
    folder structure the browser sent (needed so _find_complaint_file's "prefer a pleadings/
    subfolder" check keeps working, and so identity for incremental rescans is stable). Returns None
    if nothing safe is left (e.g. an empty or entirely "." /".." path)."""
    parts = [_sanitize_filename_segment(p) for p in raw_path.replace("\\", "/").split("/") if p not in ("", ".", "..")]
    parts = [p for p in parts if p]
    return Path(*parts) if parts else None


def _save_uploaded_case_files(job_id: str, file_storages: list) -> tuple[Path, list[Path], str]:
    """Saves an uploaded case folder's files (a Werkzeug FileStorage list, one per file picked via
    <input type="file" webkitdirectory>) into this job's own source directory, preserving each
    file's relative path so _find_complaint_file's subfolder check and incremental-rescan identity
    both keep working exactly as they did when this read directly off the reviewer's own disk.
    Returns (source_dir, list of saved .pdf paths, folder_display_name) — folder_display_name is
    just the picked folder's own top-level name (the first path segment every uploaded file shares),
    shown in the UI in place of the old absolute filesystem path, which no longer exists server-side."""
    source_dir = CASE_SOURCES_DIR / job_id
    source_dir.mkdir(parents=True, exist_ok=True)

    folder_display_name = None
    saved_pdf_paths = []
    for fs in file_storages:
        rel_path = _safe_relative_path(fs.filename or "")
        if rel_path is None:
            continue
        if folder_display_name is None and len(rel_path.parts) > 1:
            folder_display_name = rel_path.parts[0]
        dest = source_dir.joinpath(*rel_path.parts[1:]) if len(rel_path.parts) > 1 else source_dir / rel_path
        dest.parent.mkdir(parents=True, exist_ok=True)
        fs.save(dest)
        if dest.suffix.lower() == ".pdf":
            saved_pdf_paths.append(dest)

    return source_dir, sorted(saved_pdf_paths), (folder_display_name or "case files")


def _run_case_job(job_id: str, folder: Path, folder_display_name: str, detail_level: str,
                   plaintiff_name: str, priority_hint: str, defendant_names: list) -> None:
    """Runs entirely in a background thread (see /case/start). A case is one plaintiff by default —
    unlike the old "group an unknown folder of mixed patients" design this replaced, the primary
    group is always the plaintiff named when the case was started, not something inferred from
    content. The existing content/filename identity-guessing stays as a SECONDARY fallback: if the
    folder's records also reference a different, distinctly-identified individual (a co-plaintiff, a
    family case), that still gets its own separate group — it just isn't the primary experience
    anymore. Writes the manifest to disk after every meaningful step so the status page can show
    live progress via polling, and so one group's failure never loses work already done for
    another — the whole point of this being safe to leave running overnight.

    `folder` here is this job's own CASE_SOURCES_DIR/<job_id> directory (already populated by
    _save_uploaded_case_files before this thread was started), not the reviewer's original disk
    location — Flask never reads a path the reviewer typed/browsed to directly (see the module-level
    comment above CASE_SOURCES_DIR for why). `folder_display_name` is purely cosmetic — the picked
    folder's own top-level name, for the status page to show in place of an absolute path.
    """
    manifest = {
        "job_id": job_id,
        "source_dir": str(folder),
        "folder_display_name": folder_display_name,
        "plaintiff_name": plaintiff_name,
        "priority_hint": priority_hint,
        "detail_level": detail_level,  # remembered so a later rescan (see _run_case_rescan) uses
                                        # the same detail level without the reviewer re-specifying it
        "status": "scanning",
        "started_at": time.time(),
        "groups": {},
    }
    _save_batch_manifest(job_id, manifest)  # also registers it in _batch_jobs — see that function

    pdf_paths = sorted(folder.rglob("*.pdf"))
    if not pdf_paths:
        manifest["status"] = "error"
        manifest["error"] = "No PDF files found in that folder (searched recursively)."
        _save_batch_manifest(job_id, manifest)
        return

    # Phase 1: find and read the Complaint/NOI first, per the standing instruction to always start
    # a case by understanding why the suit exists before reading the medical records themselves.
    manifest["progress_text"] = "Looking for the Complaint/Notice of Intent…"
    _save_batch_manifest(job_id, manifest)

    complaint_path = _find_complaint_file(pdf_paths, folder)
    dol, facts_summary, complaint_warning = None, None, None
    if complaint_path:
        try:
            complaint_info = _extract_complaint_info(complaint_path)
            dol = complaint_info.get("dol")
            facts_summary = complaint_info.get("facts_summary")
        except Exception as e:
            complaint_warning = (
                f"Found a likely Complaint/NOI ({complaint_path.name}) but couldn't read it: {e}. "
                "Proceeding without case context — DOL and Facts will be blank."
            )
            applog.log_event(
                "complaint_extraction_failed", f"Failed to extract complaint info: {type(e).__name__}",
                level="ERROR", context={"job_id": job_id}, exc_info=sys.exc_info(),
            )
        pdf_paths = [p for p in pdf_paths if p != complaint_path]
    else:
        complaint_warning = (
            "No Complaint/Notice of Intent file was found in this case's folder (checked filenames "
            "for \"complaint\"/\"notice of intent\"/\"NOI\", preferring a 'pleadings' subfolder). "
            "DOL and Facts will be blank below — you can still review and fill these in manually on "
            "export."
        )

    manifest.update({
        "defendant_names": defendant_names,
        "dol": dol,
        "facts_summary": facts_summary,
        "complaint_source_file": complaint_path.name if complaint_path else None,
        "complaint_warning": complaint_warning,
    })
    _save_batch_manifest(job_id, manifest)

    case_context = _build_case_context(plaintiff_name, defendant_names, dol, facts_summary, priority_hint)

    # Phase 2: priority reorder — whatever the reviewer flagged as especially important (a subfolder
    # name, a file-type keyword) gets processed first, so the mid-generation live-partial-results
    # feature surfaces the important material earliest instead of in arbitrary folder-scan order.
    priority_terms = [t.strip().lower() for t in re.split(r"[,\n]", priority_hint or "") if t.strip()]
    if priority_terms:
        def is_priority(p: Path) -> bool:
            path_str = str(p).lower()
            return any(term in path_str for term in priority_terms)
        pdf_paths.sort(key=lambda p: not is_priority(p))

    # Phase 3: everything belongs to the plaintiff's own primary group by default; content/filename
    # identity-checking only pulls a file OUT into its own secondary group when it confidently finds
    # a DIFFERENT, distinct individual — never used to second-guess the plaintiff's own identity.
    primary_key = _normalize_patient_key(plaintiff_name, dob=None, fallback_id="")
    groups: dict = {primary_key: {"display_name": plaintiff_name, "files": []}}

    for i, pdf_path in enumerate(pdf_paths, start=1):
        manifest["progress_text"] = f"Scanning case files: {i} of {len(pdf_paths)}"
        _save_batch_manifest(job_id, manifest)

        name = _guess_patient_from_filename(pdf_path.name)
        dob = None
        if not name:
            try:
                text, _, _ = pdf_to_text(pdf_path)
            except Exception:
                text = ""
            identity = _extract_patient_identity(text)
            name, dob = identity["name"], identity["dob"]

        if not name or _same_person(name, plaintiff_name):
            groups[primary_key]["files"].append(pdf_path)
            continue

        # A hash of the full resolved path, not just the filename — a folder searched recursively
        # can easily have same-named files in different subfolders, and those must never collide.
        fallback_id = hashlib.md5(str(pdf_path.resolve()).encode()).hexdigest()[:10]
        key = _normalize_patient_key(name, dob, fallback_id=fallback_id)
        if key not in groups:
            # identity_name/identity_dob (not just the display label) are stored so a LATER rescan
            # (see _run_case_rescan) can match a new file against this same individual without
            # having to string-parse "Also mentions: X" back apart — fragile and easy to get wrong.
            groups[key] = {"display_name": f"Also mentions: {name}", "identity_name": name, "identity_dob": dob, "files": []}
        groups[key]["files"].append(pdf_path)

    manifest["progress_text"] = None
    for key, g in groups.items():
        manifest["groups"][key] = {
            "display_name": g["display_name"],
            "identity_name": g.get("identity_name"),
            "identity_dob": g.get("identity_dob"),
            "file_count": len(g["files"]),
            # Relative paths this group is SUPPOSED to cover, independent of db.record_processed_files
            # (which only runs once this group's own generation, incl. synthesis, is fully done — see
            # _resume_incomplete_case_jobs, which needs this to detect a group as still-incomplete
            # even in the narrow window where "status" already briefly says "done" but the files
            # haven't actually been recorded yet).
            "files": [str(p.relative_to(folder)) for p in g["files"]],
            "status": "pending",
        }
    manifest["status"] = "processing"
    _save_batch_manifest(job_id, manifest)

    # Phase 4: one chronology per group, sequentially (there's only one local model to share, so
    # concurrent Ollama calls would just queue behind each other anyway). The primary (plaintiff)
    # group gets the full case context; secondary "also mentions" groups don't, so the model isn't
    # misled into judging at_issue from a defendant relevant to a different person's records.
    for key, g in groups.items():
        session_id = f"case_{job_id}_{key}"
        # Set BEFORE generation starts, not just once done — the partial/live findings shown via
        # on_progress below need this immediately so the viewer can already fetch the source PDFs
        # for highlighting while later chunks are still generating, not just once fully finished.
        manifest["groups"][key]["status"] = "processing"
        manifest["groups"][key]["session_id"] = session_id
        _save_batch_manifest(job_id, manifest)

        def on_progress(status_text, partial_findings, key=key):
            # Default-arg trick captures the CURRENT key by value for this closure, not whatever
            # key the outer loop variable points to by the time this callback actually fires later.
            manifest["groups"][key]["progress_text"] = status_text
            if partial_findings is not None:
                # Shown live in the UI (see /case/<job>/<group>/results and the main viewer's
                # polling) — a reviewer gets a real, growing chronology to read while later chunks
                # of a large case are still generating, rather than nothing until it all finishes.
                manifest["groups"][key]["findings"] = partial_findings
            _save_batch_manifest(job_id, manifest)

        try:
            session_dir = BATCH_SESSIONS_DIR / session_id
            copied_paths = _copy_files_into_session(session_dir, g["files"])

            # Durable ids assigned BEFORE generation — they have to already be known to build this
            # call's own "### Document N —" headers. Keyed by filename (not rel_path) to match what
            # _generate_chronology's read loop looks each file up by, since _copy_files_into_session
            # preserves the original filename 1:1 (dest.name == src.name, see that function).
            rel_paths = [str(p.relative_to(folder)) for p in g["files"]]
            ids_by_rel_path = db.get_or_assign_file_ids(job_id, key, rel_paths)
            file_ids = {p.name: ids_by_rel_path[rel] for p, rel in zip(g["files"], rel_paths)}

            group_context = case_context if key == primary_key else ""
            result = _generate_chronology(
                copied_paths, detail_level, case_context=group_context, on_progress=on_progress,
                file_ids=file_ids,
            )
            manifest["groups"][key].update({
                "status": "done",
                "progress_text": None,
                "findings": result["findings"],
                "stats": result["stats"],
                "warnings": result["warnings"],
                "ocr_files": result["ocr_files"],
                "source_filenames": result["source_filenames"],
                "record_source_labels": result["record_source_labels"],
                "page_counts": result["page_counts"],
            })
            # Case-level DOL/summary synthesis — only for the primary (plaintiff) group, and only
            # after its own findings are fully in the manifest above (a synthesis failure here must
            # never lose the group's own already-successful chronology). See _synthesize_case_summary.
            if key == primary_key:
                try:
                    demographics = result["findings"].get("patient_demographics") or {}
                    synthesis = _synthesize_case_summary(
                        result["findings"].get("timeline", []), result["findings"].get("visits", []),
                        demographics, dol, facts_summary, case_context,
                    )
                    manifest["dol"] = synthesis.get("dol") or dol
                    manifest["groups"][key]["findings"]["summary"] = (
                        synthesis.get("summary") or result["findings"].get("summary")
                    )
                    # Fill gaps only — never let the synthesis pass overwrite an already-real,
                    # per-chunk-extracted value (same "first real value wins" caution as the
                    # per-chunk merge itself, just applied one layer later).
                    for field, value in (synthesis.get("patient_demographics") or {}).items():
                        if _realish(value) and not _realish(demographics.get(field)):
                            demographics[field] = value
                    manifest["groups"][key]["findings"]["patient_demographics"] = demographics
                except Exception as e:
                    applog.log_event(
                        "case_synthesis_failed", f"Case-level DOL/summary synthesis failed: {type(e).__name__}",
                        level="WARNING", context={"job_id": job_id}, exc_info=sys.exc_info(),
                    )
                _save_batch_manifest(job_id, manifest)
            # Recorded so a LATER rescan (see _run_case_rescan) knows these files are already
            # processed and only reprocesses genuinely new ones — paths relative to this job's own
            # source_dir (folder), not the copied session-dir paths, and not an absolute path either
            # (a rescan's freshly re-uploaded files never resolve to the same absolute path twice).
            db.record_processed_files(job_id, key, [str(p.relative_to(folder)) for p in g["files"]])
        except Exception as e:
            # Deliberately broad: a bad PDF, a JSON-parse hiccup, a model timeout on one group
            # must never take down the rest of an unattended overnight run.
            manifest["groups"][key]["status"] = "failed"
            manifest["groups"][key]["error"] = str(e)
            applog.log_event(
                "case_group_processing_failed", f"Group processing failed: {type(e).__name__}",
                level="ERROR", context={"job_id": job_id, "group_key": key}, exc_info=sys.exc_info(),
            )

        _save_batch_manifest(job_id, manifest)

    manifest["status"] = "done"
    manifest["finished_at"] = time.time()
    _save_batch_manifest(job_id, manifest)


def _run_case_rescan(job_id: str) -> None:
    """Runs in a background thread (see POST /case/<job_id>/rescan), AFTER the route handler has
    already saved a freshly re-uploaded copy of the reviewer's case folder into this job's
    source_dir (see _save_uploaded_case_files — files whose relative path already exists there are
    left untouched, so this is safe to call even though most of the re-upload is stuff already
    processed). Picks up where the case's initial run left off: finds files whose relative path
    within source_dir was never processed (tracked in db.py's processed_files table — see Phase 4 of
    _run_case_job, which records every file it processes), reprocesses ONLY those, and merges the
    results into the existing chronology rather than starting over. New files by relative path only
    (not content changes to an already-processed file — a deliberate scope boundary, see
    TEST_RESULTS.md).

    Also doubles as the crash-resume mechanism (see _resume_incomplete_case_jobs): when this is
    called on a job whose overall `status` ISN'T already "done" (i.e. this is a resume, not an
    ordinary reviewer-triggered rescan on an already-complete case), this function also owns
    flipping `status` to "done"/"error" once it finishes — ordinarily that's _run_case_job's job
    alone, but a resumed job never goes through that function again.
    """
    manifest = _get_batch_manifest(job_id)
    if manifest is None:
        return
    was_already_done = manifest.get("status") == "done"

    manifest["rescan_status"] = "processing"
    manifest["rescan_message"] = None
    _save_batch_manifest(job_id, manifest)

    try:
        folder = Path(manifest["source_dir"])
        pdf_paths = sorted(folder.rglob("*.pdf"))
        known_paths = db.get_processed_paths(job_id)
        complaint_name = manifest.get("complaint_source_file")
        new_paths = [
            p for p in pdf_paths
            if str(p.relative_to(folder)) not in known_paths and p.name != complaint_name
        ]

        if not new_paths:
            manifest["rescan_status"] = "done"
            manifest["rescan_message"] = "No new files found — this case's folder hasn't changed."
            manifest["rescan_finished_at"] = time.time()
            if not was_already_done:
                manifest["status"] = "done"
                manifest["finished_at"] = time.time()
            _save_batch_manifest(job_id, manifest)
            return

        plaintiff_name = manifest["plaintiff_name"]
        detail_level = manifest.get("detail_level", DEFAULT_DETAIL_LEVEL)
        primary_key = _normalize_patient_key(plaintiff_name, dob=None, fallback_id="")
        existing_groups = manifest.get("groups", {})
        case_context = _build_case_context(
            plaintiff_name, manifest.get("defendant_names") or [], manifest.get("dol"),
            manifest.get("facts_summary"), manifest.get("priority_hint", ""),
        )

        # Group the WHOLE new-files batch before creating any group (not file-by-file) — so two new
        # files that turn out to be the same new secondary individual land in ONE new group, not two.
        new_groups: dict = {}
        for pdf_path in new_paths:
            name = _guess_patient_from_filename(pdf_path.name)
            dob = None
            if not name:
                try:
                    text, _, _ = pdf_to_text(pdf_path)
                except Exception:
                    text = ""
                identity = _extract_patient_identity(text)
                name, dob = identity["name"], identity["dob"]

            if not name or _same_person(name, plaintiff_name):
                target_key = primary_key
            else:
                target_key = None
                for existing_key, existing_g in existing_groups.items():
                    if existing_key == primary_key:
                        continue
                    existing_identity = existing_g.get("identity_name")
                    if existing_identity and _same_person(name, existing_identity):
                        target_key = existing_key
                        break
                if target_key is None:
                    for new_key, new_g in new_groups.items():
                        if new_key == primary_key:
                            continue
                        if new_g.get("identity_name") and _same_person(name, new_g["identity_name"]):
                            target_key = new_key
                            break
                if target_key is None:
                    fallback_id = hashlib.md5(str(pdf_path.resolve()).encode()).hexdigest()[:10]
                    target_key = _normalize_patient_key(name, dob, fallback_id=fallback_id)

            if target_key not in new_groups:
                new_groups[target_key] = {"files": []}
                if target_key != primary_key and target_key not in existing_groups:
                    new_groups[target_key]["display_name"] = f"Also mentions: {name}"
                    new_groups[target_key]["identity_name"] = name
                    new_groups[target_key]["identity_dob"] = dob
            new_groups[target_key]["files"].append(pdf_path)

        for key, g in new_groups.items():
            is_new_group = key not in manifest["groups"]
            if is_new_group:
                manifest["groups"][key] = {
                    "display_name": g.get("display_name", plaintiff_name if key == primary_key else "Unknown"),
                    "identity_name": g.get("identity_name"),
                    "identity_dob": g.get("identity_dob"),
                    "file_count": 0,
                    "status": "pending",
                }
            session_id = manifest["groups"][key].get("session_id") or f"case_{job_id}_{key}"
            manifest["groups"][key]["session_id"] = session_id
            manifest["groups"][key]["status"] = "processing"
            manifest["groups"][key]["progress_text"] = f"Processing {len(g['files'])} new file(s)…"
            _save_batch_manifest(job_id, manifest)

            def on_progress(status_text, partial_findings, key=key):
                manifest["groups"][key]["progress_text"] = status_text
                _save_batch_manifest(job_id, manifest)

            try:
                session_dir = BATCH_SESSIONS_DIR / session_id
                copied_paths = _copy_files_into_session(session_dir, g["files"])

                if key == primary_key and not is_new_group:
                    group_context = _build_incremental_context(manifest, manifest["groups"][key])
                elif key == primary_key:
                    group_context = case_context
                else:
                    group_context = ""

                # Durable ids assigned BEFORE generation, same reasoning/keying as _run_case_job —
                # reuses whatever this job already assigned to earlier files, so a rescan's new
                # files get ids continuing from the existing max rather than colliding with them.
                new_rel_paths = [str(p.relative_to(folder)) for p in g["files"]]
                ids_by_rel_path = db.get_or_assign_file_ids(job_id, key, new_rel_paths)
                file_ids = {p.name: ids_by_rel_path[rel] for p, rel in zip(g["files"], new_rel_paths)}

                new_result = _generate_chronology(
                    copied_paths, detail_level, case_context=group_context, on_progress=on_progress,
                    file_ids=file_ids,
                )

                if is_new_group:
                    manifest["groups"][key].update({
                        "status": "done", "progress_text": None,
                        "findings": new_result["findings"], "stats": new_result["stats"],
                        "warnings": new_result["warnings"], "ocr_files": new_result["ocr_files"],
                        "source_filenames": new_result["source_filenames"],
                        "record_source_labels": new_result["record_source_labels"],
                        "page_counts": new_result["page_counts"],
                        "file_count": len(g["files"]),
                        "files": new_rel_paths,
                    })
                else:
                    merged = _merge_incremental_group_result(manifest["groups"][key], new_result)
                    manifest["groups"][key].update({
                        "status": "done", "progress_text": None,
                        "findings": merged["findings"], "stats": merged["stats"],
                        "warnings": merged["warnings"], "ocr_files": merged["ocr_files"],
                        "source_filenames": merged["source_filenames"],
                        "record_source_labels": merged["record_source_labels"],
                        "page_counts": merged["page_counts"],
                        "file_count": manifest["groups"][key].get("file_count", 0) + len(g["files"]),
                        "files": (manifest["groups"][key].get("files") or []) + new_rel_paths,
                    })

                # Re-run case-level DOL/summary synthesis so it reflects the now-updated primary
                # timeline — same reasoning and failure handling as _run_case_job's own call.
                if key == primary_key:
                    try:
                        demographics = manifest["groups"][key]["findings"].get("patient_demographics") or {}
                        synthesis = _synthesize_case_summary(
                            manifest["groups"][key]["findings"].get("timeline", []),
                            manifest["groups"][key]["findings"].get("visits", []), demographics,
                            manifest.get("dol"), manifest.get("facts_summary"), case_context,
                        )
                        manifest["dol"] = synthesis.get("dol") or manifest.get("dol")
                        manifest["groups"][key]["findings"]["summary"] = (
                            synthesis.get("summary") or manifest["groups"][key]["findings"].get("summary")
                        )
                        for field, value in (synthesis.get("patient_demographics") or {}).items():
                            if _realish(value) and not _realish(demographics.get(field)):
                                demographics[field] = value
                        manifest["groups"][key]["findings"]["patient_demographics"] = demographics
                    except Exception as e:
                        applog.log_event(
                            "case_synthesis_failed", f"Case-level DOL/summary synthesis failed: {type(e).__name__}",
                            level="WARNING", context={"job_id": job_id}, exc_info=sys.exc_info(),
                        )

                db.record_processed_files(job_id, key, [str(p.relative_to(folder)) for p in g["files"]])
            except Exception as e:
                manifest["groups"][key]["status"] = "failed"
                manifest["groups"][key]["error"] = str(e)
                applog.log_event(
                    "case_rescan_group_failed", f"Rescan group processing failed: {type(e).__name__}",
                    level="ERROR", context={"job_id": job_id, "group_key": key}, exc_info=sys.exc_info(),
                )

            _save_batch_manifest(job_id, manifest)

        manifest["rescan_status"] = "done"
        manifest["rescan_message"] = f"Processed {len(new_paths)} new file(s) across {len(new_groups)} group(s)."
        manifest["rescan_finished_at"] = time.time()
        if not was_already_done:
            manifest["status"] = "done"
            manifest["finished_at"] = time.time()
        _save_batch_manifest(job_id, manifest)
    except Exception as e:
        manifest["rescan_status"] = "done"
        manifest["rescan_message"] = f"Rescan failed: {e}"
        if not was_already_done:
            # A resumed job that still fails on retry must be visibly stuck, not silently stay
            # "processing" forever (which would otherwise re-trigger identically on every restart).
            manifest["status"] = "error"
            manifest["error"] = f"Resume after restart failed: {e}"
        applog.log_event(
            "case_rescan_failed", f"Rescan failed: {type(e).__name__}",
            level="ERROR", context={"job_id": job_id}, exc_info=sys.exc_info(),
        )
        _save_batch_manifest(job_id, manifest)


def _resume_incomplete_case_jobs() -> None:
    """Called once at server startup, before app.run(). A case job (or a "check for new files"
    rescan) runs in a background thread tied to this process — if the machine loses power or the
    app is force-quit mid-run, that thread is gone forever, but the manifest itself survives
    (SQLite, written after every meaningful step — see _save_batch_manifest). Without this, such a
    job is silently stuck at "processing" forever with no way to finish short of starting an
    entirely new case from scratch (the exact complaint from a real 50-file case: a laptop died
    halfway through).

    Resumes by reusing the SAME codepath as "Check for new files" (_run_case_rescan) — it already
    knows, per group, exactly which files were never fully processed (db.get_processed_paths is
    only populated once a group's own chronology generation fully completes — see
    db.record_processed_files at the end of both _run_case_job's and _run_case_rescan's per-group
    loops), so a group that had already finished before the crash is left completely untouched, and
    one that was mid-way through gets redone from that group's own start — not resumed mid-chunk
    (nothing in this pipeline checkpoints at that finer a grain), but still far better than losing
    the whole case.
    """
    for manifest in db.list_case_manifests():
        job_id = manifest["job_id"]
        status = manifest.get("status")

        if status == "done" and manifest.get("rescan_status") == "processing":
            # An interrupted "Check for new files" rescan on an otherwise-complete case — safe to
            # just re-run directly: rescan's own on_progress never touches a group's findings
            # mid-flight (only progress_text does), so nothing partial needs clearing first.
            applog.log_event(
                "case_rescan_resumed", "Resuming interrupted rescan after restart",
                context={"job_id": job_id},
            )
            threading.Thread(target=_run_case_rescan, args=(job_id,), daemon=True).start()
            continue

        if status not in ("scanning", "processing"):
            continue  # "done" (no stuck rescan) / "error" — nothing to resume

        if not manifest.get("groups"):
            # Crashed before Phase 3 (grouping) ever completed — nothing group-specific was
            # persisted yet, so a full fresh run (complaint extraction included) is simplest and
            # correct; there's no partial per-group state to preserve at this early a stage.
            applog.log_event(
                "case_job_resumed", "Resuming interrupted case job (pre-grouping) after restart",
                context={"job_id": job_id},
            )
            threading.Thread(
                target=_run_case_job, args=(
                    job_id, Path(manifest["source_dir"]), manifest.get("folder_display_name", ""),
                    manifest.get("detail_level", DEFAULT_DETAIL_LEVEL), manifest["plaintiff_name"],
                    manifest.get("priority_hint", ""), manifest.get("defendant_names") or [],
                ), daemon=True,
            ).start()
            continue

        # Past grouping — one or more groups may have been mid-generation when the crash happened.
        # A group's own "status" field alone isn't trustworthy here: status flips to "done" (see
        # _run_case_job/_run_case_rescan) BEFORE db.record_processed_files runs (that write happens
        # last, after the case-level synthesis pass) — a crash in that narrow window would leave a
        # group showing "done" with genuinely complete, correct findings, yet none of its files
        # actually recorded as processed. Cross-check against the real source of truth instead.
        processed_paths = db.get_processed_paths(job_id)
        stuck = []
        for key, g in manifest["groups"].items():
            group_files = set(g.get("files") or [])
            if g.get("status") == "done" and (not group_files or group_files.issubset(processed_paths)):
                continue  # genuinely complete and recorded — leave entirely untouched
            if g.get("status") == "done" and group_files and not group_files.issubset(processed_paths):
                # Findings are already complete and correct — only the bookkeeping write was
                # interrupted. Finish that cheaply rather than wastefully regenerating good results.
                db.record_processed_files(job_id, key, sorted(group_files))
                continue
            stuck.append(key)

        # Clear a genuinely-stuck group's transient generation state (findings/stats/etc, all
        # possibly a PARTIAL snapshot from the interrupted run's own on_progress callback) before
        # handing off to rescan, so its merge-with-existing-findings starts from empty rather than
        # double-counting whatever had already been saved right before the crash. file_count resets
        # to 0 for the same reason — rescan's merge branch ADDS the reprocessed file count on top of
        # whatever's already there (see _run_case_rescan), which was already the group's full
        # intended count from Phase 3, so leaving it as-is would double it.
        for key in stuck:
            manifest["groups"][key].update({
                "status": "pending", "progress_text": None, "findings": {}, "stats": {},
                "warnings": [], "ocr_files": [], "source_filenames": [],
                "record_source_labels": {}, "page_counts": {}, "file_count": 0,
            })
        manifest["status"] = "processing"
        _save_batch_manifest(job_id, manifest)
        applog.log_event(
            "case_job_resumed", f"Resuming interrupted case job after restart ({len(stuck)} group(s))",
            context={"job_id": job_id},
        )
        threading.Thread(target=_run_case_rescan, args=(job_id,), daemon=True).start()


@app.route("/case")
def case_page():
    return render_template("case.html")


@app.route("/case/start", methods=["POST"])
def case_start():
    """Case folders are uploaded from the browser's own native folder picker (<input type="file"
    webkitdirectory> in case.html), not read off a path typed/browsed to on whichever machine Flask
    happens to be running on — see the CASE_SOURCES_DIR comment near the top of this file for why
    that older design was replaced. `files` here is every PDF (and everything else) the picker found
    inside the chosen folder, each carrying its path relative to that folder as its filename."""
    plaintiff_name = request.form.get("plaintiff_name", "").strip()
    defendant_names = [n.strip() for n in re.split(r"[,\n]", request.form.get("defendant_names", "")) if n.strip()]
    priority_hint = request.form.get("priority_hint", "").strip()
    detail_level = request.form.get("detail_level", DEFAULT_DETAIL_LEVEL)
    if detail_level not in DETAIL_LEVELS:
        detail_level = DEFAULT_DETAIL_LEVEL

    if not plaintiff_name:
        return jsonify({"error": "Please provide the plaintiff's name"}), 400
    file_storages = request.files.getlist("files")
    if not file_storages:
        return jsonify({"error": "Please select the case's folder"}), 400

    job_id = uuid.uuid4().hex
    source_dir, saved_pdf_paths, folder_display_name = _save_uploaded_case_files(job_id, file_storages)
    if not saved_pdf_paths:
        return jsonify({"error": "No PDF files found in that folder (searched subfolders too)."}), 400

    thread = threading.Thread(
        target=_run_case_job,
        args=(job_id, source_dir, folder_display_name, detail_level, plaintiff_name, priority_hint, defendant_names),
        daemon=True,
    )
    thread.start()
    applog.log_event(
        "case_started", "New case started",
        context={"job_id": job_id, "file_count": len(saved_pdf_paths), "detail_level": detail_level},
    )
    return jsonify({"job_id": job_id})


def _compute_review_status(job_id: str, group_key: str, group: dict, decided_map: dict) -> dict:
    """Derives {"total", "decided"} for one group — never stored, always computed fresh from
    review_actions/added_facts so it can't drift out of sync with them. `decided_map` is a single
    job-wide db.count_decided() result, passed in so callers iterating every group in a job don't
    re-run that query per group. Clamped because delete_added_fact doesn't cascade-delete a matching
    review_actions row today, so a fact deleted after being approved could otherwise push decided
    above total."""
    total = _count_ai_findings(group.get("findings") or {}) + len(db.load_added_facts(job_id, group_key))
    decided = min(decided_map.get(group_key, 0), total)
    return {"total": total, "decided": decided}


@app.route("/case/status/<job_id>")
def case_status(job_id):
    manifest = _get_batch_manifest(job_id)
    if manifest is None:
        return jsonify({"error": "Unknown job id"}), 404
    result = dict(manifest)
    exported_map = db.get_exported_map(job_id)
    decided_map = db.count_decided(job_id)
    result["groups"] = {
        key: {
            **g,
            "review_progress": _compute_review_status(job_id, key, g, decided_map),
            "exported_at": exported_map.get(key),
        }
        for key, g in manifest.get("groups", {}).items()
    }
    return jsonify(result)


@app.route("/case/<job_id>/rescan", methods=["POST"])
def case_rescan(job_id):
    """Kicks off incremental reprocessing: only files added to this case's folder since it was
    last processed (see _run_case_rescan) — manually triggered (a button on the case status page or
    the review page, not automatic background polling), so this never surprises the reviewer with
    unexpected model usage while they're not looking.

    Expects a fresh multipart upload of the SAME folder, re-picked via the same native folder
    picker used to start the case (see case_start) — there's no server-side path to re-scan on its
    own anymore, since the whole point of the upload-based redesign is that Flask never needs to
    read an arbitrary path off whichever disk it's running on. Re-saving into this job's existing
    source_dir is harmless for files it's seen before (same relative path, just overwritten) and is
    how genuinely new files actually arrive."""
    manifest = _get_batch_manifest(job_id)
    if manifest is None:
        return jsonify({"error": "Unknown job id"}), 404
    if manifest.get("status") != "done":
        return jsonify({"error": "This case is still processing its first pass — wait for it to finish first."}), 409
    if manifest.get("rescan_status") == "processing":
        return jsonify({"error": "A rescan is already running for this case."}), 409
    file_storages = request.files.getlist("files")
    if not file_storages:
        return jsonify({"error": "Please re-select the case's folder to check for new files."}), 400

    _save_uploaded_case_files(job_id, file_storages)

    thread = threading.Thread(target=_run_case_rescan, args=(job_id,), daemon=True)
    thread.start()
    applog.log_event("case_rescan_started", "Rescan started", context={"job_id": job_id})
    return jsonify({"status": "started"})


@app.route("/case/jobs")
def case_jobs_list():
    """Lists all known case jobs, most recent first — lets the status page be reopened (even in
    a fresh browser tab, even after this server restarted) without needing the job id memorized,
    since a job you started before going to bed is exactly the kind of thing you'd otherwise lose
    track of by morning. Also backs the dashboard's case list (see templates/dashboard.html) — the
    richer fields here (defendant names, DOL, document counts, the primary group's own key/
    readiness) are everything already sitting in the manifest, just projected out, so the dashboard
    needs no new persistence of its own."""
    jobs = []
    for manifest in db.list_case_manifests():
        groups = manifest.get("groups", {})
        primary_key = _normalize_patient_key(manifest.get("plaintiff_name"), dob=None, fallback_id="")
        primary_group = groups.get(primary_key, {})
        failed_group_count = sum(1 for g in groups.values() if g.get("status") == "failed")
        # Only the primary group's review/export status is surfaced here — it's the only group a
        # dashboard row links to (see caseAction() in dashboard.js); secondary groups are visible via
        # the case status page, which pulls the full per-group breakdown from /case/status/<job_id>.
        job_id = manifest["job_id"]
        primary_review_progress = None
        primary_exported_at = None
        if primary_group:
            decided_map = db.count_decided(job_id)
            primary_review_progress = _compute_review_status(job_id, primary_key, primary_group, decided_map)
            primary_exported_at = db.get_exported_map(job_id).get(primary_key)
        jobs.append({
            "job_id": job_id,
            "folder_display_name": manifest.get("folder_display_name"),
            "plaintiff_name": manifest.get("plaintiff_name"),
            "defendant_names": manifest.get("defendant_names") or [],
            "dol": manifest.get("dol"),
            "status": manifest["status"],
            "started_at": manifest.get("started_at"),
            "finished_at": manifest.get("finished_at"),
            "total_documents": sum(g.get("file_count", 0) for g in groups.values()),
            "group_count": len(groups),
            "failed_group_count": failed_group_count,
            "primary_group_key": primary_key,
            "primary_group_ready": bool(primary_group.get("findings")),
            "primary_group_review_progress": primary_review_progress,
            "primary_group_exported_at": primary_exported_at,
        })
    return jsonify(jobs)


@app.route("/case/<job_id>/<group_key>/results")
def case_group_results(job_id, group_key):
    manifest = _get_batch_manifest(job_id)
    if manifest is None:
        return jsonify({"error": "Unknown job"}), 404
    group = manifest.get("groups", {}).get(group_key)
    if not group:
        # A genuine, real bug found via live testing: `manifest["groups"]` stays completely empty
        # for as long as Phase 1-3 takes (finding/reading the Complaint/NOI, then grouping
        # documents by identity) — often well over a minute on a large case — before the primary
        # (plaintiff) group entry even gets created. A reviewer who opens the review link the
        # moment a case starts (the normal, expected way to use this) will have their very FIRST
        # poll land in that window and get a 404 here. The frontend's poll loop (see
        # fetchAndRenderCaseResult in app.js) treats ANY `error` field as fatal and never schedules
        # another poll — so a 404 at exactly this moment silently killed the page forever, with no
        # further updates ever, even hours after the job actually finished. As long as the overall
        # job hasn't reached a terminal state, "this group doesn't exist yet" is a completely
        # normal transient condition, not a real error — respond exactly like the "still working,
        # keep polling" shape below so the frontend's existing retry logic just works. Only once
        # the job itself is done/errored and the group STILL never showed up is this a genuine,
        # permanent "unknown group" — e.g., a stale/mistyped link to a group that never existed.
        if manifest.get("status") not in ("done", "error"):
            return jsonify({
                "ready": False,
                "status": manifest.get("status"),
                "progress_text": manifest.get("progress_text"),
                "case": None,
            })
        return jsonify({"error": "Unknown group"}), 404
    # Serve partial results while still processing (see the on_progress callback in
    # _run_case_job) — a reviewer gets a real, growing chronology to read immediately rather
    # than waiting for a potentially long multi-chunk analysis to fully finish. Only a group with
    # no findings at all yet (still on its very first chunk, or still just reading documents) has
    # nothing to show.
    # Case-level metadata (plaintiff/defendant/DOL/facts) only makes sense attached to the PRIMARY
    # (plaintiff) group's results — a secondary "also mentions" group gets an empty case block
    # (see _run_case_job, which likewise skips case_context for those groups).
    primary_key = _normalize_patient_key(manifest.get("plaintiff_name"), dob=None, fallback_id="")
    case_meta = {
        "plaintiff_name": manifest.get("plaintiff_name"),
        "defendant_names": manifest.get("defendant_names", []),
        "dol": manifest.get("dol"),
        "facts_summary": manifest.get("facts_summary"),
        "complaint_source_file": manifest.get("complaint_source_file"),
        "complaint_warning": manifest.get("complaint_warning"),
        "source_filenames": group.get("source_filenames", []),
    } if group_key == primary_key else None
    if not group.get("findings"):
        # Not an error — this is the normal, expected state while the first chunk is still being
        # read/generated, and the frontend polls this every few seconds. Returning 200 here (rather
        # than 404) keeps the browser console clean during ordinary polling; a real 404 is reserved
        # for a genuinely unknown job/group above.
        return jsonify({
            "ready": False,
            "status": group.get("status"),
            "progress_text": group.get("progress_text"),
            "case": case_meta,
        })
    return jsonify({
        "ready": True,
        "session_id": group["session_id"],
        "findings": group["findings"],
        "stats": group.get("stats"),
        "warnings": group.get("warnings", []),
        "ocr_files": group.get("ocr_files", []),
        "status": group.get("status"),
        "progress_text": group.get("progress_text"),
        "case": case_meta,
        "record_source_labels": group.get("record_source_labels", {}),
        "page_counts": group.get("page_counts", {}),
        # Durable reviewer state (approve/reject/edits/added facts) — lets a genuinely fresh
        # browser tab (no sessionStorage for this session yet, e.g. reopening a case from a month
        # ago) pick up exactly where a previous review session left off. See db.py's
        # review_actions/added_facts tables; the frontend pre-seeds sessionStorage from these on
        # first load, then sessionStorage stays the synchronous source of truth as before.
        "review_state": db.load_review_state(job_id, group_key),
        "added_facts": db.load_added_facts(job_id, group_key),
        "metadata_overrides": db.load_metadata_overrides(job_id, group_key),
        "review_progress": _compute_review_status(job_id, group_key, group, db.count_decided(job_id)),
        "exported_at": db.get_exported_map(job_id).get(group_key),
        "rescan_status": manifest.get("rescan_status"),
        "rescan_message": manifest.get("rescan_message"),
    })


# --- Durable reviewer-state endpoints (Case Mode only — see db.py) -----------------------------
# Deliberately thin: each just validates the job/group exist, then calls straight into db.py.
# app.js calls these as fire-and-forget writes alongside its existing sessionStorage saves (see
# saveReviewStatus/saveFindingEdits/saveSummaryEdit/saveAddedFacts) — sessionStorage stays the
# synchronous source of truth for rendering; these are a durability backstop only.

def _require_case_group(job_id, group_key):
    """Shared validation for the review-state write endpoints below. Returns an error Response to
    return immediately, or None if the job/group are valid."""
    manifest = _get_batch_manifest(job_id)
    if manifest is None:
        return jsonify({"error": "Unknown job"}), 404
    if group_key not in manifest.get("groups", {}):
        return jsonify({"error": "Unknown group"}), 404
    return None


@app.route("/case/<job_id>/<group_key>/review", methods=["POST"])
def case_save_review(job_id, group_key):
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    data = request.get_json(force=True) or {}
    finding_id = data.get("finding_id")
    if not finding_id:
        return jsonify({"error": "finding_id required"}), 400
    db.save_review_action(job_id, group_key, finding_id, status=data.get("status"))
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/edit", methods=["POST"])
def case_save_edit(job_id, group_key):
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    data = request.get_json(force=True) or {}
    finding_id = data.get("finding_id")
    if not finding_id:
        return jsonify({"error": "finding_id required"}), 400
    db.save_review_action(job_id, group_key, finding_id, edited_text=data.get("text"))
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/summary_edit", methods=["POST"])
def case_save_summary_edit(job_id, group_key):
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    data = request.get_json(force=True) or {}
    db.save_review_action(job_id, group_key, "__summary__", kind="summary", edited_text=data.get("text"))
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/metadata_edit", methods=["POST"])
def case_save_metadata_edit(job_id, group_key):
    """Durable backstop for the preview modal's click-to-edit case-content values (DOL, Facts,
    demographics, defendant names) — same fire-and-forget-alongside-sessionStorage pattern as
    /edit and /summary_edit above, just keyed by an arbitrary field_key instead of a finding id."""
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    data = request.get_json(force=True) or {}
    field_key = data.get("field_key")
    if not field_key:
        return jsonify({"error": "field_key required"}), 400
    db.save_metadata_override(job_id, group_key, field_key, data.get("value") or "")
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/facts", methods=["POST"])
def case_save_added_fact(job_id, group_key):
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    data = request.get_json(force=True) or {}
    fact_id = data.get("fact_id")
    if not fact_id:
        return jsonify({"error": "fact_id required"}), 400
    db.save_added_fact(job_id, group_key, str(fact_id), data.get("payload") or {})
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/facts/<fact_id>", methods=["DELETE"])
def case_delete_added_fact(job_id, group_key, fact_id):
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    db.delete_added_fact(job_id, group_key, fact_id)
    return jsonify({"ok": True})


@app.route("/case/<job_id>/<group_key>/export", methods=["POST"])
def case_mark_exported(job_id, group_key):
    """Records that this group's chronology was exported, for the "Exported" review-status badge
    (see _compute_review_status/dashboard.js/case.js). Deliberately separate from /export/docx
    itself, which stays the stateless "dumb renderer" its own docstring describes — this is called
    by app.js as a fire-and-forget follow-up once a docx download actually succeeds."""
    err = _require_case_group(job_id, group_key)
    if err:
        return err
    db.mark_exported(job_id, group_key)
    return jsonify({"ok": True})


@app.route("/analyze", methods=["POST"])
def analyze():
    if "zipfile" not in request.files or not request.files["zipfile"].filename:
        return jsonify({"error": "No file uploaded"}), 400

    upload = request.files["zipfile"]
    if not upload.filename.lower().endswith(".zip"):
        return jsonify({"error": "Please upload a .zip file"}), 400

    detail_level = request.form.get("detail_level", DEFAULT_DETAIL_LEVEL)
    if detail_level not in DETAIL_LEVELS:
        detail_level = DEFAULT_DETAIL_LEVEL

    # Read the upload into memory now, before handing off to the streaming generator below —
    # keeps this simple regardless of how Flask/Werkzeug manage the request context during streaming.
    zip_bytes = upload.read()

    def generate():
        _clear_old_sessions()
        session_id = uuid.uuid4().hex
        session_dir = SESSIONS_DIR / session_id
        session_dir.mkdir(parents=True, exist_ok=True)

        zip_path = session_dir / "_upload.zip"
        zip_path.write_bytes(zip_bytes)

        try:
            pdf_paths = safe_extract_pdfs(zip_path, session_dir)
        except zipfile.BadZipFile:
            yield _event({"stage": "error", "message": "That doesn't look like a valid zip file"})
            return
        finally:
            zip_path.unlink(missing_ok=True)  # keep the extracted PDFs, not the raw zip

        if not pdf_paths:
            yield _event({"stage": "error", "message": "No PDF files found inside that zip"})
            return

        pdf_paths = sorted(pdf_paths, key=lambda p: p.name)
        total_files = len(pdf_paths)
        yield _event({"stage": "start", "total_files": total_files, "session_id": session_id})

        records_parts = []
        warnings = []
        ocr_files = []  # filenames where at least one page needed the OCR fallback — the browser
                         # viewer can't highlight passages on these (the PDF has no text layer for
                         # pdf.js to search), so the frontend uses this to show an accurate message
                         # instead of the generic "couldn't locate the quote" one.
        filename_to_text = {}  # for _resolve_bates below — the exact per-file text the model saw
        file_id_to_filename = {}  # for _resolve_file_ids — converts the model's file_id back to a real filename
        page_counts = {}  # for the Record Sources legend's "N page(s)" — a hard fact, never model-derived
        total_pages = 0

        for i, pdf_path in enumerate(pdf_paths, start=1):
            try:
                text, num_pages, num_ocr_pages = pdf_to_text(pdf_path)
            except Exception as e:
                warnings.append(f"Could not read {pdf_path.name}: {e}")
                applog.log_event(
                    "document_extraction_failed", f"Failed to extract text from a document: {type(e).__name__}",
                    level="ERROR", context={"document_index": i, "document_count": total_files},
                    exc_info=sys.exc_info(),
                )
                yield _event({
                    "stage": "extract", "index": i, "total": total_files,
                    "filename": pdf_path.name, "pages": 0, "ok": False,
                })
                continue

            total_pages += num_pages
            clean_text = text.strip()
            ok = True

            if num_ocr_pages:
                ocr_files.append(pdf_path.name)
                warnings.append(
                    f"'{pdf_path.name}': {num_ocr_pages} of {num_pages} page(s) had no usable text "
                    "layer (likely scanned/faxed) and were read via OCR instead — double-check "
                    "those findings against the source image, since OCR is less reliable than a "
                    "native text layer."
                )

            if not clean_text:
                warnings.append(
                    f"'{pdf_path.name}': no extractable text found even after OCR "
                    f"({num_pages} page(s)) — this file was excluded from the analysis."
                )
                ok = False
                yield _event({
                    "stage": "extract", "index": i, "total": total_files,
                    "filename": pdf_path.name, "pages": num_pages, "ok": False,
                })
                continue

            records_parts.append(f"### Document {i} — {pdf_path.name}\n\n{clean_text}")
            filename_to_text[pdf_path.name] = clean_text
            file_id_to_filename[i] = pdf_path.name
            page_counts[pdf_path.name] = num_pages
            yield _event({
                "stage": "extract", "index": i, "total": total_files,
                "filename": pdf_path.name, "pages": num_pages, "ok": ok,
            })

        if not records_parts:
            yield _event({
                "stage": "error",
                "message": "No extractable text found in any uploaded PDF.",
                "warnings": warnings,
            })
            return

        records_text = "\n\n---\n\n".join(records_parts)

        truncated = False
        if len(records_text) > MAX_INPUT_CHARS:
            records_text = records_text[:MAX_INPUT_CHARS]
            truncated = True
            warnings.append(
                "Combined document text was longer than this model's context window and was "
                "truncated to fit — not all content was analyzed. Try fewer/shorter documents per "
                "run, or raise LAWFIRMAGENT_NUM_CTX (uses more memory)."
            )

        yield _event({
            "stage": "extract_done",
            "documents_processed": len(records_parts),
            "total_pages": total_pages,
            "input_truncated": truncated,
        })

        prompt = PROMPT_TEMPLATE.format(
            records=records_text,
            detail_instructions=DETAIL_LEVELS[detail_level],
            case_context="No case context was provided for this run.",
        )
        payload = {
            "model": MODEL,
            "prompt": prompt,
            "stream": True,
            # Computed from records_text AFTER truncation above — a document truncated off entirely
            # correctly has no valid id here, since it has no usable content in this call either way.
            "format": _response_schema_for_ids(_valid_file_ids_in_text(records_text)),
            "options": {"num_ctx": NUM_CTX},
        }

        start = time.time()
        full_response = ""
        tokens_seen = 0
        eval_count = 0
        eval_duration_s = 0

        try:
            with requests.post(OLLAMA_URL, json=payload, stream=True, timeout=900) as resp:
                resp.raise_for_status()
                for line in resp.iter_lines():
                    if not line:
                        continue
                    chunk = json.loads(line)
                    full_response += chunk.get("response", "")

                    if chunk.get("done"):
                        eval_count = chunk.get("eval_count", tokens_seen)
                        eval_duration_s = chunk.get("eval_duration", 0) / 1e9
                        break

                    tokens_seen += 1
                    yield _event({
                        "stage": "generating",
                        "tokens_so_far": tokens_seen,
                        "elapsed": round(time.time() - start, 1),
                    })
        except requests.exceptions.RequestException as e:
            yield _event({
                "stage": "error",
                "message": f"Could not reach the local model server ({e}). Is Ollama running? "
                           "Check with: launchctl list | grep lawfirmagent",
            })
            return

        wall_time = time.time() - start

        try:
            findings = json.loads(full_response)
        except json.JSONDecodeError:
            yield _event({
                "stage": "error",
                "message": "The model's response wasn't valid JSON this run — this can happen "
                           "occasionally with schema-constrained output. Try again.",
            })
            return

        _normalize_dates(findings)
        _resolve_file_ids(findings, file_id_to_filename)
        _reconcile_source_files(findings, filename_to_text)
        _resolve_bates(findings, filename_to_text)
        _dedupe_all_with_multi_source(findings)
        _assign_finding_ids(findings)
        record_source_labels = _resolve_record_sources(findings, list(filename_to_text.keys()))
        _verify_demographics_with_grep(findings, "\n\n".join(filename_to_text.values()))

        yield _event({
            "stage": "done",
            "session_id": session_id,
            "findings": findings,
            "stats": {
                "documents_processed": len(records_parts),
                "total_pages": total_pages,
                "wall_time_seconds": round(wall_time, 1),
                "output_tokens": eval_count,
                "tokens_per_second": round(eval_count / eval_duration_s, 1) if eval_duration_s else None,
                "input_truncated": truncated,
                "model": MODEL,
                "detail_level": detail_level,
            },
            "warnings": warnings,
            "ocr_files": ocr_files,
            "record_source_labels": record_source_labels,
            "page_counts": page_counts,
            "source_filenames": list(filename_to_text.keys()),
        })

    resp = Response(stream_with_context(generate()), mimetype="application/x-ndjson")
    resp.headers["Cache-Control"] = "no-cache"
    resp.headers["X-Accel-Buffering"] = "no"  # no-op locally, harmless if ever put behind a proxy
    resp.direct_passthrough = True
    return resp


if __name__ == "__main__":
    # Inside a container this must be 0.0.0.0 instead: Docker's port-publishing forwards from the
    # HOST's 127.0.0.1 to the container's own network namespace, and a container-internal bind to
    # 127.0.0.1 isn't reachable via that forwarding. The container is still not exposed beyond the
    # host, because the compose file publishes this port as `127.0.0.1:5050:5050` on the host side
    # — see docker-compose.yml. (_BIND_HOST itself is set at module load — see its definition above.)
    _resume_incomplete_case_jobs()
    app.run(host=_BIND_HOST, port=PORT, debug=False, threaded=True)
